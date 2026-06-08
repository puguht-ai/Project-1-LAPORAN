from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# ==================== PAGE 1 ====================
add_title("Screening Berita Ekonomi Terpanas — Fraksi Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("Fraksi Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 1 dari 4", 9)
doc.add_paragraph()
add_title("SCREENING BERITA EKONOMI TERPANAS", 14)
add_title("Fiskal  ·  Moneter  ·  APBN 2026", 12)
add_centered("Periode: 25–30 Mei 2026     |     Perspektif: Anggota DPR Fraksi Penyeimbang     |     Disusun: 30 Mei 2026", 10)
doc.add_paragraph()

add_normal("Periode", 10, True)
add_normal("25–30 Mei 2026", 10)
add_normal("Topik Dominan", 10, True)
add_normal("Rupiah · BI Rate · Likuiditas Perbankan · Ekspor SDA", 10)
add_normal("Skala Tekanan", 10, True)
add_normal("🔴 Sangat Tinggi", 10)
doc.add_paragraph()

add_normal("⚖ Posisi Fraksi Penyeimbang: Fungsi penyeimbang bukan sekadar menolak — melainkan memastikan kebijakan pemerintah diuji dengan standar data, logika, dan kepentingan publik tertinggi. Seluruh kritik di bawah ini bersifat evidence-based, bukan partisan.", 10, True)
doc.add_paragraph()

# BERITA 1
add_normal("Berita 1   Moneter · Fiskal", 11, True)
add_normal("Rupiah Nyaris Rp17.900 per Dolar AS — Menkeu Sri Mulyani: \"Tak Masuk Akal\"", 10, True)
add_normal("25–30 Mei 2026    |  Moneter  Fiskal", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Rupiah terus melemah dan mendekati level psikologis Rp17.900/USD pada periode 25–30 Mei 2026 — mengkhawatirkan pasar dan pelaku ekonomi.", 10)
add_normal("▸ Rupiah kini berada di bawah Dollar Singapura — alarm baru bagi ekonomi Indonesia yang menunjukkan pelemahan mata uang domestik secara signifikan dibandingkan negara tetangga.", 10)
add_normal("▸ Menkeu Sri Mulyani menyatakan kondisi rupiah \"tak masuk akal\" — pengakuan resmi bahwa situasi nilai tukar sudah melampaui batas kewajaran.", 10)
add_normal("▸ Pemicu utama: Oil shock akibat konflik AS-Iran mendorong harga minyak Brent ke USD101-102/barel, capital outflow dari pasar saham Indonesia, dan kekhawatiran investor terhadap defisit fiskal yang melebar.", 10)
add_normal("▸ Mata uang Asia babak belur secara keseluruhan akibat oil shock — rupiah menjadi salah satu mata uang dengan pelemahan terburuk di kawasan.", 10)
add_normal("▸ BI memperketat aturan pembelian dolar AS untuk membantu stabilisasi nilai tukar — langkah kebijakan untuk mengurangi tekanan pada mata uang domestik.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Pernyataan Menkeu \"tak masuk akal\" adalah pengakuan telat yang seharusnya memicu respons kebijakan konkret, bukan hanya ekspresi keheranan.", 10)
add_normal("● Rupiah di bawah Dollar Singapura adalah aib ekonomi nasional yang tidak boleh dianggap enteng.", 10)
add_normal("● Oil shock yang masih berlangsung dan potensi eskalasi konflik AS-Iran menjadi ancaman berkelanjutan.", 10)
add_normal("● BI memperketat aturan beli dolar adalah langkah korektif yang baik — tetapi ini adalah instrumen taktis.", 10)
doc.add_paragraph()

# BERITA 2
add_normal("Berita 2   Moneter · Fiskal", 11, True)
add_normal("BI Rate Kembali Dinaikkan — Rupiah Tertekan, Ruang Stimulus Pertumbuhan Makin Sempit", 10, True)
add_normal("25–30 Mei 2026    |  Moneter  Fiskal", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Bank Indonesia (BI) kembali menaikkan BI-Rate dalam periode Mei 2026 sebagai respons terhadap tekanan inflasi dan pelemahan rupiah.", 10)
add_normal("▸ Kenaikan BI-Rate dilakukan untuk menarik kembali modal asing dan menstabilkan rupiah — namun berdampak pada biaya pinjaman domestik yang semakin mahal.", 10)
add_normal("▸ LPS mempertahankan tingkat bunga penjaminan simpanan meskipun BI-Rate naik — menunjukkan kebijakan untuk menjaga stabilitas sistem perbankan nasional.", 10)
add_normal("▸ Ruang stimulus pertumbuhan ekonomi semakin terbatas karena kebijakan moneter yang harus memprioritaskan stabilitas nilai tukar daripada ekspansi kredit.", 10)
add_normal("▸ Ekonom memperingatkan bahwa kenaikan BI-Rate berkepanjangan akan menekan sektor riil — kredit usaha menjadi lebih mahal dan investasi melambat.", 10)
add_normal("▸ Kondisi ini menciptakan dilema kebijakan: tahan rupiah (bunga naik, pertumbuhan turun) atau biarkan rupiah melemah (beban utang valas naik, daya beli turun).", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Kenaikan BI-Rate adalah pilihan yang menyakitkan di antara dua pilihan buruk.", 10)
add_normal("● Setiap 25 basis poin kenaikan BI-Rate menambah beban bunga utang negara yang besar.", 10)
add_normal("● LPS mempertahankan bunga penjaminan adalah langkah stabilitas yang tepat.", 10)
add_normal("● Dilema moneter ini adalah konsekuensi dari ketidaksinkronan kebijakan fiskal-moneter.", 10)
doc.add_page_break()

# ==================== PAGE 2 ====================
add_title("Screening Berita Ekonomi Terpanas — Fraksi Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("Fraksi Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 2 dari 4", 9)
doc.add_paragraph()

# BERITA 3
add_normal("Berita 3   Moneter · Fiskal", 11, True)
add_normal("Empat Bank Jumbo Hadapi Badai Likuiditas — BBCA, BMRI, BBRI, BBNI Tertekan", 10, True)
add_normal("25–30 Mei 2026    |  Moneter  Perbankan", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Empat bank terbesar Indonesia — BBCA, BMRI, BBRI, dan BBNI — menghadapi tantangan likuiditas yang signifikan pada periode Mei 2026.", 10)
add_normal("▸ Tekanan likuiditas ini terjadi seiring dengan kenaikan BI-Rate dan upaya stabilisasi nilai tukar yang memerlukan likuiditas dalam dolar AS.", 10)
add_normal("▸ Kondisi ini menunjukkan tekanan pada sektor perbankan nasional yang harus menyeimbangkan antara pertumbuhan kredit dan kebutuhan likuiditas valas.", 10)
add_normal("▸ Bank Mandiri (BMRI) mencatatkan pertumbuhan kredit sebesar 18,5% YoY hingga April 2026, namun menghadapi tantangan likuiditas di tengah kondisi pasar yang ketat.", 10)
add_normal("▸ Maybank Indonesia mencatatkan laba sebelum pajak sebesar Rp397 miliar di kuartal I 2026.", 10)
add_normal("▸ Kebutuhan likuiditas valas untuk stabilisasi rupiah menambah tekanan pada likuiditas perbankan domestik.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Empat bank jumbo adalah tulang punggung sistem keuangan Indonesia.", 10)
add_normal("● Pertumbuhan kredit 18,5% adalah pencapaian yang patut diapresiasi.", 10)
add_normal("● Kebutuhan likuiditas valas untuk intervensi pasar menambah kompleksitas.", 10)
add_normal("● Stabilisasi rupiah bukan hanya tanggung jawab BI dan pemerintah.", 10)
doc.add_paragraph()

# BERITA 4
add_normal("Berita 4   Fiskal · Ekspor", 11, True)
add_normal("Prabowo Wajibkan Ekspor SDA Lewat BUMN DSI — Sawit dan Batubara Kena Dampak", 10, True)
add_normal("25–30 Mei 2026    |  Fiskal  Ekspor", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Presiden Prabowo memutuskan kewajiban ekspor Sumber Daya Alam (SDA) melalui BUMN yang akan dibentuk.", 10)
add_normal("▸ Ekspor komoditas strategis seperti kelapa sawit dan batubara akan terdampak oleh kebijakan baru ini.", 10)
add_normal("▸ Kebijakan ini merupakan bagian dari upaya pemerintah untuk meningkatkan nilai tambah dan mengoptimalkan penerimaan negara.", 10)
add_normal("▸ Dana Hasil Ekspor (DHE) SDA dan Badan Ekspor akan diimplementasikan mulai Juni 2026.", 10)
add_normal("▸ Tujuannya adalah memastikan bahwa ekspor SDA memberikan manfaat maksimal bagi perekonomian nasional.", 10)
add_normal("▸ Kebijakan ini berpotensi meningkatkan penerimaan negara, namun juga menimbulkan kekhawatiran di kalangan pelaku usaha.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Kewajiban ekspor melalui BUMN DSI adalah kebijakan yang berani.", 10)
add_normal("● DHE SDA yang dimulai Juni 2026 perlu dimonitor ketat dampaknya.", 10)
add_normal("● Potensi kenaikan nilai tambah dari pengolahan dalam negeri harus dimaksimalkan.", 10)
add_normal("● Kekhawatiran pelaku usaha mengenai kompleksitas perlu direspons dengan petunjuk teknis yang jelas.", 10)
doc.add_paragraph()

# BERITA 5
add_normal("Berita 5   Fiskal · Pasar Modal", 11, True)
add_normal("Asing Net Sell Rp8,54 Triliar di Akhir Mei 2026 — Investor Asing Lanjutkan Capital Outflow", 10, True)
add_normal("25–30 Mei 2026    |  Fiskal  Pasar Modal", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Investor asing melanjutkan capital outflow dengan mencatatkan net sell sebesar Rp8,54 triliun di akhir Mei 2026.", 10)
add_normal("▸ Capital outflow ini menyeret saham-saham blue chip Indonesia.", 10)
add_normal("▸ Tren net sell ini melanjutkan pola dari periode MSCI rebalancing sebelumnya.", 10)
add_normal("▸ Tekanan pada IHSG berlanjut seiring dengan capital outflow yang terus berlangsung.", 10)
add_normal("▸ Kondisi ini diperparah oleh pelemahan rupiah dan ketidakpastian kebijakan domestik.", 10)
add_normal("▸ Meskipun demikian, beberapa sektor tertentu masih menarik minat investor asing.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Net sell Rp8,54 triliun dalam sebulan adalah angka yang signifikan.", 10)
add_normal("● Capital outflow yang berkelanjutan mencerminkan kekhawatiran struktural.", 10)
add_normal("● Blue chip yang diterjang menunjukkan bahwa bahkan perusahaan terbaik tidak kebal terhadap sentimen negatif.", 10)
add_normal("● Diperlukan kebijakan yang lebih meyakinkan dan konsisten.", 10)
doc.add_page_break()

# ==================== PAGE 3 ====================
add_title("Screening Berita Ekonomi Terpanas — Fraksi Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("Fraksi Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 3 dari 4", 9)
doc.add_paragraph()

add_title("SUMBER & CATATAN KAKI", 12)

add_normal("¹ Kontan, 25–30 Mei 2026 — \"Rupiah Nyaris Rp17.900 per Dollar AS\"", 9)
add_normal("² Kontan, 25–30 Mei 2026 — \"Rupiah Kalah dari Dollar Singapura\"", 9)
add_normal("³ Kontan, 25–30 Mei 2026 — \"Mata Uang Asia Babak Belur Akibat Oil Shock\"", 9)
add_normal("⁴ Bisnis.com, 25–30 Mei 2026 — \"Nilai Tukar Rupiah Masih Loyo, BI Perketat Aturan Beli Dolar AS\"", 9)
add_normal("¹ Bisnis.com, 25–30 Mei 2026 — \"LPS Pertahankan Tingkat Bunga Penjaminan Usai BI Rate Naik\"", 9)
add_normal("² Bisnis.com, 25–30 Mei 2026 — \"Bank Mandiri Cetak Laba Rp18,1 Triliun, Kredit Tumbuh 18,5%\"", 9)
add_normal("³ Bisnis.com, 25–30 Mei 2026 — \"Maybank Indonesia Catat Laba Sebelum Pajak Rp397 Miliar\"", 9)
add_normal("¹ Kontan, 25–30 Mei 2026 — \"Prabowo Wajibkan Ekspor SDA Lewat BUMN DSI\"", 9)
add_normal("² Kontan, 25–30 Mei 2026 — \"DHE SDA dan Badan Ekspor Diimplementasikan Mulai Juni\"", 9)
add_normal("¹ Kontan, 25–30 Mei 2026 — \"Asing Net Sell Rp8,54 Triliar di Akhir Mei 2026\"", 9)

doc.add_page_break()

# ==================== PAGE 4 ====================
add_title("Screening Berita Ekonomi Terpanas — François Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("François Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 4 dari 4", 9)
doc.add_paragraph()

add_title("RINGKASAN EKSEKUTIF", 12)

add_normal("Periode pelaporan: 25–30 Mei 2026", 10)
add_normal("Topik dominan: Rupiah · BI Rate · Likuiditas Perbankan · Ekspor SDA", 10)
add_normal("Skala tekanan: 🔴 Sangat Tinggi", 10)
doc.add_paragraph()

add_normal("Lima berita terpanas periode ini:", 10, True)
add_normal("1. Rupiah nyaris Rp17.900/USD — Menkeu akui \"tak masuk akal\"", 10)
add_normal("2. BI Rate kembali dinaikkan — ruang stimulus menyempit", 10)
add_normal("3. Empat bank jumbo hadapi badai likuiditas", 10)
add_normal("4. Prabowo wajibkan ekspor SDA lewat BUMN DSI", 10)
add_normal("5. Asing net sell Rp8,54 triliun di akhir Mei", 10)
doc.add_paragraph()

add_normal("Kesimpulan:", 10, True)
add_normal("Ekonomi Indonesia menghadapi tekanan multi-dimensi: nilai tukar melemah, kebijakan moneter ketat, sektor perbankan tertekan, dan capital outflow berlanjut. Diperlukan solusi struktural, bukan hanya instrumen taktis.", 10)

# Save
output_path = r"C:\Users\USER\Screening_Berita_Ekonomi_25-30Mei2026.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")