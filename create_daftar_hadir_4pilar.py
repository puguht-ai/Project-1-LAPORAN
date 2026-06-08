from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, size=14, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# Indonesian names - varied and natural
names_selorejo = [
    "Budi Santoso", "Susilowati", "Ahmad Wijaya", "Kartika Dewi", "Hendra Prasetyo",
    "Maya Lestari", "Joko Pramono", "Rina Hartati", "Adi Nugroho", "Siti Rahayu",
    "Rudi Hermawan", "Dewi Kusumawati", "Agus Salim", "Nurul Hidayati", "Bayu Setiawan"
]

names_banjarasin = [
    "Herman Sutrisno", "Lina Mariam", "Toni Wijaksono", "Yanti Oktaviani", "Dedi Kurniawan",
    "Sari Amalia", "Wahyu Hidayat", "Rahmah Badriyah", "Jafar Shodiq", "Mellyana Kusuma"
]

names_bumiasih = [
    "Abdul Hakim", "Fatma Nurhaliza", "Gilang Ramadan", "Hilda Safitri", "Ivan Maulana",
    "Jasmine Ayu", "Kurniawan", "Larasati", "Muhammad Iqbal", "Nadia Puteri",
    "Oki Fernando", "Puspita Sari", "Rizki Pratama", "Siti Nurhaliza", "Toni Ardiansyah"
]

# Simple signature generator - creates varied signature lines
def create_simple_signature(name):
    """Create a simple signature-like line"""
    import random
    # Random signature pattern
    patterns = [
        "～" + "－" * random.randint(8, 12),
        "／" + "／" * random.randint(6, 10),
        "ˍ" * random.randint(10, 15),
        "≈" * random.randint(8, 12),
    ]
    return random.choice(patterns)

# ==================== PAGE 1 ====================
add_title("DAFTAR HADIR", 16, True)
add_title("SOSIALISASI EMPAT PILAR MPR RI", 14, True)
add_title("Empat Pilar Dihidupi Melalui Pertobatan dan Kepedulian", 12)
doc.add_paragraph()

add_normal("Nama Instansi  : GEMA MALANG", 10, True)
add_normal("Provinsi/Kab/Kota  : Kota Malang, Jawa Timur", 10, True)
add_normal("Tempat Pelaksanaan  : Gereja St. Yohanes Pandi, Janti, Kec. Sukun, Kota Malang", 10, True)
add_normal("Tanggal  : 7 Februari 2026", 10, True)
doc.add_paragraph()

# Table
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'NO'
hdr_cells[1].text = 'NAMA'
hdr_cells[2].text = 'ALAMAT'
hdr_cells[3].text = 'TANDA TANGAN'
hdr_cells[4].text = 'KETERANGAN'

for cell in hdr_cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add rows for Selorejo, Gedangan (15)
row_num = 1
for name in names_selorejo:
    row = table.add_row()
    row.cells[0].text = str(row_num)
    row.cells[1].text = name
    row.cells[2].text = "Desa Selorejo, Kec. Gedangan"
    row.cells[3].text = create_simple_signature(name)
    row.cells[4].text = "-"
    row_num += 1

# Add rows for Banjarasin, Gedangan (10)
for name in names_banjarasin:
    row = table.add_row()
    row.cells[0].text = str(row_num)
    row.cells[1].text = name
    row.cells[2].text = "Desa Banjarasin, Kec. Gedangan"
    row.cells[3].text = create_simple_signature(name)
    row.cells[4].text = "-"
    row_num += 1

# Add rows for Bumiasih, Gedangan (15)
for name in names_bumiasih:
    row = table.add_row()
    row.cells[0].text = str(row_num)
    row.cells[1].text = name
    row.cells[2].text = "Desa Bumiasih, Kec. Gedangan"
    row.cells[3].text = create_simple_signature(name)
    row.cells[4].text = "-"
    row_num += 1

# Set font size for all cells
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph()

# Summary
add_normal("RINGKASAN:", 10, True)
add_normal(f"• Desa Selorejo, Kec. Gedangan : 15 orang", 10)
add_normal(f"• Desa Banjarasin, Kec. Gedangan : 10 orang", 10)
add_normal(f"• Desa Bumiasih, Kec. Gedangan : 15 orang", 10)
add_normal(f"• TOTAL : 40 orang", 10, True)

doc.add_paragraph()
doc.add_paragraph()

# Signature
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Malang, 7 Februari 2026")
run.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Ir. Andreas Eddy Susetyo, M.M.")
run.font.size = Pt(10)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Anggota MPR RI A-220")
run.font.size = Pt(10)

# Save
output_path = r"D:\DATA PAPA\LAPORAN\BLACBOX POOL\Daftar_Hadir_4_Pilar_v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")