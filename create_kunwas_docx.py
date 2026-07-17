from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper functions
def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_centered(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_photo(path, width=Inches(4)):
    try:
        doc.add_picture(path, width=width)
    except Exception as e:
        print(f"Error adding photo: {e}")

# ==================== PAGE 1: SAMPUL ====================
add_title("DEWAN PERWAKILAN RAKYAT REPUBLIK INDONESIA", 16)
doc.add_paragraph()
add_title("LAPORAN KUNJUNGAN KERJA PENGAWASAN PERSEORANGAN", 18)
add_title("KE - 2", 18)
add_title("MASA SIDANG V TAHUN SIDANG 2025-2026", 14)

doc.add_paragraph()
doc.add_paragraph()

add_centered("Nama                        : Ir. Andreas Eddy Susetyo, M.M.")
add_centered("Nomor Anggota      : A-220")
add_centered("Komisi                       : XI")
add_centered("Fraksi                        : PDI-Perjuangan")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 5 - 7 Juni 2026")

doc.add_page_break()

# ==================== I. PENDAHULUAN ====================
add_title("I. PENDAHULUAN", 14)

pendahuluan_text = """Dalam rangka pelaksanaan tugas konstitusional yakni menjalankan fungsi pengawasan anggota Dewan Perwakilan Rakyat Republik Indonesia terhadap pelaksanaan kebijakan pemerintah, maka dilaksanakan Kunjungan Kerja Pengawasan Perseorangan Ke-2 yang berlangsung pada tanggal 5 hingga 7 Juni 2026 di wilayah Malang Raya, pada Masa Sidang V Tahun Sidang 2025-2026. Kunjungan kali ini difokuskan secara khusus untuk memantau, mengawasi, dan mengevaluasi situasi di masyarakat terkait kondisi perekonomian yang sedang mengalami tekanan berat dan tidak berada dalam kondisi yang baik-baik saja. Hal ini ditandai oleh memburuknya indikator makroekonomi secara berkelanjutan, baik di tingkat nasional maupun lokal, yang secara langsung menggerus ketahanan ekonomi rumah tangga dan kelangsungan usaha masyarakat.

Secara nasional, lanskap makroekonomi pada pertengahan tahun 2026 menunjukkan tren yang mengkhawatirkan. Tingkat inflasi tahunan, meskipun secara agregat masih dalam target, mengalami tekanan inflasi inti yang persisten di kisaran 3,5% hingga 4,2%, yang terutama didorong oleh volatilitas harga pangan dan energi. Pelemahan nilai tukar Rupiah terhadap Dolar AS yang terus berlangsung telah meningkatkan biaya impor bahan baku, mesin, dan komoditas strategis, yang pada akhirnya diteruskan kepada konsumen dalam bentuk kenaikan harga barang jadi. Di sisi moneter, kebijakan suku bunga acuan yang tetap tinggi untuk meredam inflasi justru berdampak pada mahalnya biaya kredit, sehingga menghambat ekspansi sektor riil. Pertumbuhan Produk Domestik Bruto (PDB) nasional pada kuartal II-2026 diproyeksikan melambat dan stagnan di kisaran 4,8% hingga 5,0%, angka yang berada di bawah potensi pertumbuhan optimal Indonesia dan mencerminkan lesunya permintaan domestik.

Dampak dari kondisi makro yang memburuk ini terasa sangat nyata dan sistemik di tingkat lokal Malang Raya. Sebagai daerah yang ekonominya sangat bergantung pada sektor perdagangan, jasa, pariwisata, dan UMKM, Malang Raya mengalami kontraksi permintaan yang tajam. Data daerah menunjukkan adanya peningkatan Tingkat Pengangguran Terbuka (TPT), terutama di kalangan pemuda dan pekerja sektor informal. Stagnasi daya beli masyarakat menjadi masalah kronis; pendapatan riil rumah tangga tidak lagi mampu mengimbangi laju kenaikan biaya hidup, khususnya untuk kebutuhan pokok, pendidikan, dan kesehatan. Sektor UMKM, yang merupakan tulang punggung perekonomian daerah, menghadapi ancaman serius akibat margin keuntungan yang tergerus oleh kenaikan biaya operasional, logistik, dan energi, sementara kemampuan menaikkan harga jual sangat terbatas karena daya beli konsumen yang sudah jenuh.

Dalam konteks krisis multidimensi inilah, kunjungan kerja pengawasan ini dirancang secara sistematis dan komprehensif. Tidak hanya berfokus pada harga komoditas semata, kunjungan ini memperluas cakupan pemantauan untuk melihat bagaimana tekanan ekonomi memengaruhi berbagai lapisan sosial: mulai dari kelangsungan hidup pelaku usaha, peran institusi keagamaan sebagai jaring pengaman sosial alternatif, beban ganda yang dipikul oleh perempuan dan ibu rumah tangga, ancaman kepunahan ekosistem seni-budaya akibat minimnya apresiasi ekonomi, hingga perjuangan harian masyarakat akar rumput di sektor informal.

Melalui serangkaian dialog terbuka, peninjauan lapangan, dan Rapat Dengar Pendapat Umum (RDPU) dengan berbagai pemangku kepentingan, kunjungan ini bertujuan untuk menyerap aspirasi secara langsung dan tidak tersaring. Hasil dari pengawasan ini diharapkan dapat menjadi landasan empiris yang kuat untuk merumuskan rekomendasi kebijakan yang tepat sasaran, responsif, dan berpihak pada kepentingan rakyat, serta mendesak pemerintah pusat dan daerah untuk segera mengambil langkah-langkah afirmatif guna memulihkan daya beli dan melindungi kelompok masyarakat yang paling rentan terdampak oleh badai ekonomi ini."""

for paragraph in pendahuluan_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_page_break()

# ==================== II. DASAR HUKUM ====================
add_title("II. DASAR HUKUM", 14)

add_normal("Fungsi pengawasan merupakan salah satu dari tiga fungsi utama DPR RI sebagaimana diamanatkan oleh konstitusi dan peraturan perundang-undangan di Indonesia. Dasar hukum utama yang melandasi fungsi ini adalah sebagai berikut:")

doc.add_paragraph()
add_normal("Undang-Undang Dasar Negara Republik Indonesia Tahun 1945 (UUD 1945)", bold=True)
add_bullet("Pasal 20A ayat (1): Menyatakan secara eksplisit bahwa DPR memiliki fungsi legislasi, fungsi anggaran, dan fungsi pengawasan.")
add_bullet("Pasal 20A ayat (2): Memberikan hak-hak kepada DPR dalam menjalankan fungsinya, yaitu hak interpelasi, hak angket, dan hak menyatakan pendapat.")
add_bullet("Pasal 20A ayat (3): Memberikan hak kepada setiap anggota DPR untuk mengajukan pertanyaan, menyampaikan usul dan pendapat, serta hak imunitas.")

doc.add_paragraph()
add_normal("Undang-Undang Nomor 17 Tahun 2014 tentang MPR, DPR, DPD, dan DPRD (UU MD3)", bold=True)
add_bullet("Sebagaimana telah beberapa kali diubah, terakhir dengan Undang-Undang Nomor 13 Tahun 2019 (Perubahan Ketiga). UU ini mengatur secara rinci mengenai tata cara pelaksanaan fungsi pengawasan, tugas, wewenang, serta alat kelengkapan dewan yang menjalankan pengawasan tersebut.")

doc.add_paragraph()
add_normal("Peraturan DPR RI Nomor 1 Tahun 2025 tentang Tata Tertib", bold=True)
add_bullet("Merupakan revisi terbaru dari Peraturan DPR RI Nomor 1 Tahun 2020. Peraturan ini mengatur mekanisme internal DPR dalam menjalankan fungsi pengawasan, termasuk penambahan kewenangan evaluatif (Pasal 228A) terhadap pejabat negara yang ditetapkan melalui rapat paripurna, serta prosedur rapat-rapat dan kunjungan kerja.")

doc.add_page_break()

# ==================== III. TUJUAN ====================
add_title("III. TUJUAN", 14)

add_normal("Fungsi pengawasan dijalankan untuk memastikan penyelenggaraan negara berjalan sesuai dengan koridor hukum dan kepentingan rakyat. Berikut adalah tujuan utama dari fungsi pengawasan DPR RI:")

doc.add_paragraph()
add_bullet("Menjamin Pelaksanaan Undang-Undang: Memastikan bahwa setiap kebijakan dan tindakan yang diambil oleh pemerintah (eksekutif) tidak menyimpang dari Undang-Undang yang telah ditetapkan bersama.")
add_bullet("Menjaga Akuntabilitas Anggaran: Mengawasi penggunaan Anggaran Pendapatan dan Belanja Negara (APBN) agar digunakan secara efisien, efektif, transparan, dan bertanggung jawab demi kemakmuran rakyat.")
add_bullet("Meningkatkan Kinerja Birokrasi: Mendorong pemerintah untuk senantiasa meningkatkan kualitas pelayanan publik dan kinerja instansi pemerintah melalui evaluasi berkala dalam rapat-rapat kerja.")
add_bullet("Check and Balances: Menjalankan prinsip saling mengawasi dan mengimbangi antarlembaga negara agar tidak terjadi pemusatan kekuasaan atau penyalahgunaan wewenang (abuse of power) oleh pihak eksekutif.")
add_bullet("Menyerap dan Menindaklanjuti Aspirasi Rakyat: Menjadi jembatan bagi masyarakat untuk menyampaikan keluhan atau permasalahan terkait kebijakan pemerintah, yang kemudian diklarifikasi kepada mitra kerja terkait.")

doc.add_page_break()

# ==================== IV. MEKANISME PELAKSANAAN PENGAWASAN ====================
add_title("IV. MEKANISME PELAKSANAAN PENGAWASAN", 14)

add_bullet("Rapat Kerja (Raker): Pertemuan antara komisi di DPR dengan menteri atau kepala lembaga untuk membahas kebijakan atau permasalahan tertentu.")
add_bullet("Rapat Dengar Pendapat (RDP): Rapat untuk mendengarkan penjelasan dari pejabat pemerintah di bawah tingkat menteri atau instansi terkait.")
add_bullet("Rapat Dengar Pendapat Umum (RDPU): Pertemuan dengan masyarakat, organisasi, atau pakar untuk mendapatkan masukan terkait isu tertentu.")
add_bullet("Kunjungan Kerja (Kunker): Peninjauan langsung ke lapangan oleh anggota DPR untuk melihat implementasi kebijakan atau program pemerintah di daerah.")
add_bullet("Panitia Kerja (Panja) dan Panitia Khusus (Pansus): Pembentukan tim khusus untuk mendalami permasalahan spesifik yang membutuhkan perhatian mendalam.")

doc.add_page_break()

# ==================== V. RANGKAIAN KEGIATAN ====================
add_title("V. RANGKAIAN KEGIATAN", 14)

photo_folder = r"D:\DATA PAPA\LAPORAN\KUNWAS\KUNWAS 2"

# KEGIATAN 1
add_title("KEGIATAN 1", 12)
keg1_text = """Pertemuan dialog dengan kalangan pelaku usaha dan bisnis dari berbagai sektor di Kota Malang berlangsung intens dan penuh keprihatinan. Para pengusaha menyampaikan bahwa kondisi makroekonomi yang memburuk secara langsung memukul sektor riil. Tingginya suku bunga kredit perbankan, fluktuasi nilai tukar, dan stagnasi daya beli masyarakat menyebabkan omzet penjualan turun drastis hingga 20-30 persen dibandingkan tahun sebelumnya. Banyak pelaku usaha, terutama di sektor kuliner, ritel, dan jasa, terpaksa melakukan efisiensi dengan mengurangi jam operasional atau merumahkan sebagian karyawan. Para pengusaha mendesak pemerintah pusat dan daerah untuk segera memberikan insentif fiskal, kemudahan akses kredit lunak (KUR) dengan bunga subsidi, serta deregulasi perizinan yang menghambat."""
add_normal(keg1_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P1.HEIC"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P2.HEIC"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 2
add_title("KEGIATAN 2", 12)
keg2_text = """Kunjungan dan dialog dengan umat Katolik di Paroki Hati Kudus Yesus Kayutangan, Kota Malang, menyoroti peran vital institusi keagamaan sebagai jaring pengaman sosial di tengah krisis ekonomi. Para pengurus paroki dan umat menyampaikan bahwa gereja saat ini tidak hanya berfungsi sebagai tempat ibadah, tetapi juga menjadi tumpuan bagi masyarakat yang terdampak PHK dan kesulitan ekonomi melalui program dapur umum dan beasiswa pendidikan. Namun, mereka mengakui bahwa donasi dan kolekte juga mengalami penurunan seiring dengan menurunnya kemampuan ekonomi umat. Umat Katolik mengapresiasi kehadiran wakil rakyat dan berharap ada kebijakan afirmatif dari pemerintah, seperti subsidi silang untuk sekolah berbasis yayasan dan program kemitraan ekonomi yang melibatkan komunitas gereja."""
add_normal(keg2_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P3.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P4.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 3
add_title("KEGIATAN 3", 12)
keg3_text = """Menghadiri acara keagamaan Umat Hindu se-Malang Raya di Pantai Sendang, Desa Tambakrejo, Kecamatan Sumbermanjing Wetan, memberikan perspektif unik tentang irisan antara kehidupan spiritual dan ekonomi lokal. Acara ini, yang biasanya menjadi magnet pariwisata religi, mengalami penurunan partisipasi akibat mahalnya biaya transportasi dan akomodasi. Para pemuka agama dan panitia mengungkapkan kekhawatiran bahwa tradisi dan ritual keagamaan terancam tergerus oleh himpitan ekonomi masyarakat. Diharapkan pemerintah daerah memberikan dukungan logistik dan promosi terhadap event keagamaan berbasis budaya, bukan hanya sebagai wujud toleransi, tetapi juga sebagai strategi stimulasi ekonomi mikro untuk menggerakkan roda perekonomian di wilayah pesisir selatan Malang."""
add_normal(keg3_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P5.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P6.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 4
add_title("KEGIATAN 4", 12)
keg4_text = """Bertemu dan menyerap aspirasi dari kalangan ibu-ibu dan perempuan di Kecamatan Sumbermanjing Wetan, Kabupaten Malang, mengungkap realitas pahit di garis depan pertahanan ekonomi keluarga. Para ibu mengungkapkan bahwa mereka adalah pihak yang paling merasakan dampak inflasi, terutama kenaikan harga beras, minyak goreng, dan kebutuhan sekolah anak. Dengan pendapatan suami yang stagnan atau bahkan berkurang, para perempuan dituntut untuk menjadi lebih kreatif, seperti beralih ke usaha mikro rumahan atau kerajinan tangan, namun terkendala oleh minimnya modal dan akses pemasaran. Aspirasi utama yang disampaikan adalah permintaan akan program pelatihan keterampilan yang terintegrasi dengan jaminan pemasaran, akses ke Kredit Usaha Rakyat (KUR), serta jaminan bantuan sosial (seperti PKH dan BPJS), demi menjaga ketahanan pangan dan pendidikan anak-anak mereka."""
add_normal(keg4_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P7.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P8.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 5
add_title("KEGIATAN 5", 12)
keg5_text = """Berdialog dengan beberapa tokoh budaya Malang Raya di Kepanjen, Kabupaten Malang, menyoroti ancaman serius terhadap keberlangsungan ekosistem seni dan budaya di tengah resesi ekonomi. Para seniman, dalang, dan penggiat budaya menyampaikan bahwa pesanan pentas seni dan acara adat menurun drastis karena masyarakat dan pemerintah daerah memprioritaskan pengeluaran untuk kebutuhan pokok. Hal ini menyebabkan banyak pelaku budaya beralih profesi atau hidup dalam kemiskinan struktural, yang berpotensi menghilangkan warisan budaya lokal. Tokoh budaya menekankan bahwa sektor kreatif seharusnya tidak dipandang sebagai beban, melainkan sebagai aset ekonomi yang dapat menciptakan lapangan kerja dan menarik wisatawan."""
add_normal(keg5_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P9.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P10.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 6
add_title("KEGIATAN 6", 12)
keg6_text = """Bertemu dan berdialog langsung dengan masyarakat di jalanan sekitar Pasar di Desa Sitiarjo, Kecamatan Sumbermanjing Wetan, memberikan gambaran paling otentik tentang perjuangan ekonomi kelas akar rumput. Para pedagang kaki lima, buruh harian lepas, dan pengemudi ojek online menyampaikan bahwa pendapatan harian mereka semakin tidak menentu dan tidak lagi cukup untuk memenuhi kebutuhan pokok satu hari penuh. Kenaikan harga sembako yang terus terjadi membuat mereka harus berutang ke warung atau rentenir untuk bertahan hidup. Mereka juga mengeluhkan minimnya lapangan kerja formal di wilayah selatan Malang, yang memaksa banyak pemuda menjadi pengangguran atau bekerja di sektor informal dengan upah di bawah standar."""
add_normal(keg6_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P11.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "P12.JPG"), width=Inches(5.5))
doc.add_page_break()

# ==================== VI. PENUTUP ====================
add_title("VI. PENUTUP", 14)

penutup_text = """Rangkaian Kunjungan Kerja Pengawasan Perseorangan Ke-2 yang telah berlangsung selama tiga hari di wilayah Malang Raya telah memberikan potret yang sangat jelas, mendalam, dan mengkhawatirkan mengenai dampak nyata memburuknya kondisi makroekonomi terhadap kehidupan masyarakat di tingkat akar rumput. Kunjungan ini secara tegas mengonfirmasi bahwa tekanan ekonomi tidak lagi hanya berupa statistik di atas kertas, melainkan telah bertransformasi menjadi krisis keberlangsungan hidup bagi jutaan warga.

Sintesis dari enam kegiatan lapangan menunjukkan adanya korelasi langsung antara kebijakan makro dengan penderitaan mikro. Di sektor riil (Kegiatan 1), pelaku usaha mengalami kontraksi pendapatan hingga 30%, memaksa efisiensi yang berujung pada pemutusan hubungan kerja. Ketika jaring pengaman sosial formal dari pemerintah dirasa kurang memadai atau lambat, institusi keagamaan seperti Paroki Hati Kudus Yesus (Kegiatan 2) terpaksa mengambil alih peran sebagai penyangga sosial terakhir, meski sumber daya mereka sendiri semakin menipis. Hal ini mengindikasikan adanya celah besar dalam sistem perlindungan sosial nasional yang harus segera diperbaiki.

Selain itu, temuan di Pantai Sendang (Kegiatan 3) dan Kepanjen (Kegiatan 5) mengungkap bahwa resesi ekonomi juga merupakan ancaman serius bagi keberlanjutan budaya dan pariwisata lokal. Penurunan partisipasi dalam acara keagamaan dan seni bukan hanya kerugian spiritual atau kultural, tetapi juga pukulan telak bagi ekonomi mikro warga sekitar yang menggantungkan hidup pada ekosistem tersebut. Sementara itu, dialog dengan para ibu di Sumbermanjing Wetan (Kegiatan 4) dan masyarakat informal di Pasar Desa Sitiarjo (Kegiatan 6) menjadi pengingat paling keras bahwa perempuan dan pekerja harian adalah pihak yang paling menanggung beban inflasi. Mereka terjebak dalam siklus utang dan pilihan sulit antara makanan, kesehatan, dan pendidikan anak, akibat pendapatan yang tidak menentu dan minimnya akses terhadap lapangan kerja formal.

Berdasarkan temuan-temuan lapangan yang otentik ini, beberapa rekomendasi kebijakan mendesak harus segera diperjuangkan. Pertama, pemerintah perlu memberikan insentif fiskal targeted dan kemudahan akses Kredit Usaha Rakyat (KUR) dengan bunga subsidi khusus untuk UMKM yang terdampak, disertai penundaan atau keringanan retribusi daerah. Kedua, program bantuan sosial (seperti PKH dan BPNT) harus diperluas cakupannya, diperbaiki mekanisme penyalurannya agar tepat sasaran tanpa pungutan liar, dan diprioritaskan bagi rumah tangga yang dipimpin oleh perempuan atau pekerja informal. Ketiga, pemerintah daerah harus mengalokasikan anggaran afirmatif untuk menghidupkan kembali event budaya dan keagamaan sebagai stimulus ekonomi lokal, sekaligus memperbanyak program Padat Karya Tunai di wilayah selatan Malang untuk menyerap tenaga kerja informal.

Setiap aspirasi, keluhan, dan harapan yang telah diserap selama kunjungan ini akan saya bawa dan perjuangkan dengan sungguh-sungguh dalam setiap forum rapat kerja, rapat dengar pendapat, dan sidang paripurna di DPR RI. Sebagai Anggota Komisi XI dari Fraksi PDI-Perjuangan, saya berkomitmen untuk terus mengawal agar fungsi check and balances berjalan efektif, mendesak pemerintah agar tidak hanya fokus pada stabilitas makro di atas kertas, tetapi juga pada kesejahteraan riil rakyat di lapangan. Ekonomi rakyat harus dilindungi, daya beli harus dipulihkan, dan keadilan sosial harus menjadi ruh dari setiap kebijakan yang dihasilkan oleh negara."""

for paragraph in penutup_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 8 Juni 2026")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("No Anggota: A-220")

# Save document
output_path = r"D:\DATA PAPA\LAPORAN\KUNWAS\KUNWAS 2\FINAL_LAPORAN KUNJUNGAN PENGWASAN KE-2_ 5_7 jUNI 2026_v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
