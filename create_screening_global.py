from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

# ==================== PAGE 1 ====================
add_title("Screening Berita Ekonomi Global — François Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("François Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 1 dari 4", 9)
doc.add_paragraph()
add_title("SCREENING BERITA EKONOMI GLOBAL", 14)
add_title("Impact: Fiskal  ·  Moneter  ·  APBN 2026", 12)
add_centered("Periode: 25–30 Mei 2026     |     Perspektif: Anggota DPR François Penyeimbang     |     Disusun: 30 Mei 2026", 10)
doc.add_paragraph()

add_normal("Periode", 10, True)
add_normal("25–30 Mei 2026", 10)
add_normal("Topik Dominan", 10, True)
add_normal("Oil Prices · Fed Rates · Global Inflation · Trade · Markets", 10)
add_normal("Skala Tekanan", 10, True)
add_normal("🔴 Sangat Tinggi", 10)
doc.add_paragraph()

add_normal("⚖ Posisi François Penyeimbang: Fungsi penyeimbang bukan sekadar menolak — melainkan memastikan kebijakan pemerintah diuji dengan standar data, logika, dan kepentingan publik tertinggi. Seluruh kritik di bawah ini bersifat evidence-based, bukan partisan.", 10, True)
doc.add_paragraph()

# BERITA 1
add_normal("Berita 1   Energi · Global", 11, True)
add_normal("Oil Shock Berlanjut — Brent Cetak Loss Terbesar Sejak 2020, Household AS Bayar Extra $450", 10, True)
add_normal("25–30 Mei 2026    |  Energi  Global", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Harga minyak Brent mencatat kerugian bulanan terbesar sejak 2020 — penurunan 40%+ dari puncak, namun masih elevated di atas asumsi APBN Indonesia.", 10)
add_normal("▸ Selat Hormuz tetap menjadi titik kritis utama — gangguan jalur pengiriman minyak dunia tetap menjadi ancaman.", 10)
add_normal("▸ Rata-rata rumah tangga AS membayar tambahan $450 untuk bensin dan energi akibat konflik AS-Iran.", 10)
add_normal("▸ Meskipun ada harapan kesepakatan AS-Iran, ketidakpastian geopolitik tetap tinggi.", 10)
add_normal("▸ ICP Indonesia yang diasumsikan USD70/barel sudah jebol signifikan — potensi beban tambahan ratusan triliun.", 10)
add_normal("▸ Jika harga minyak melonjak ke USD130/barel (skenario eskalasi), tekanan APBN bisa mencapai ratusan triliun tambahan.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Oil shock global langsung berdampak pada ICP Indonesia — asumsi USD70/barel sudah tidak realistis.", 10)
add_normal("● Setiap USD1/barel kenaikan dari asumsi menambah beban Rp6-7 triliun — hitung sendiri dampaknya.", 10)
add_normal("● Household AS额外 $450 = dampak langsung pada daya beli global = mempengaruhi ekspor Indonesia.", 10)
add_normal("● Skenario USD130/barel bukan fiksi ilmiah — pemerintah perlu memiliki rencana mitigasi.", 10)
doc.add_paragraph()

# BERITA 2
add_normal("Berita 2   Moneter · Global", 11, True)
add_normal("Fed Holds Rates — Pasar Ekspektasikan Penurunan, Dollar AS Tetap Kuat Tekan Rupiah", 10, True)
add_normal("25–30 Mei 2026    |  Moneter  Global", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Federal Reserve AS mempertahankan suku bunga di level tinggi — kebijakan moneter ketat berlanjut.", 10)
add_normal("▸ Pasar mengekspektasikan penurunan suku bunga di semester II-2026 — namun Fed tetap hawkish.", 10)
add_normal("▸ Dollar AS tetap kuat secara global — mata uang Asia termasuk rupiah terus tertekan.", 10)
add_normal("▸ Kebijakan Fed yang dipertahankan tinggi membuat biaya pinjaman Indonesia tetap mahal.", 10)
add_normal("▸ Selisih suku bunga Indonesia-AS yang lebar menarik carry trade keluar dari Indonesia.", 10)
add_normal("▸ Kondisi ini memperparah tekanan pada rupiah dan capital outflow dari pasar obligasi Indonesia.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Fed yang tetap hawkish membuat ruang kebijakan BI semakin terbatas — harus memilih antara pertumbuhan atau stabilitas.", 10)
add_normal("● Dollar AS kuat = rupiah lemah = beban utang valas Indonesia membengkak = defisit fiskal membesar.", 10)
add_normal("● Carry trade keluar = tekanan pada rupiah =BI harus naikkan bunga = pertumbuhan tertekan.", 10)
add_normal("● Tanpa perbaikan fundamental, siklus ini akan berlanjut — solusi struktural diperlukan.", 10)
doc.add_page_break()

# ==================== PAGE 2 ====================
add_title("Screening Berita Ekonomi Global — François Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("François Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 2 dari 4", 9)
doc.add_paragraph()

# BERITA 3
add_normal("Berita 3   Inflasi · Global", 11, True)
add_normal("Inflasi Global Bertahan Tinggi — UK Energy Bills Melonjak 13%, \"Double Scar\" Ekonomi", 10, True)
add_normal("25–30 Mei 2026    |  Inflasi  Global", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Tagihan energi rumah tangga Inggris melonjak 13% — dampak langsung dari oil shock.", 10)
add_normal("▸ Fenomena \"double scar\" (luka ganda) dari inflasi masa lalu dan shock geopolitik saat ini.", 10)
add_normal("▸ Inflasi global yang bertahan tinggi membatasi ruang kebijakan moneter di seluruh dunia.", 10)
add_normal("▸ Indonesia tidak kebal dari dinamika global — inflasi domestik tetap terjaga tinggi.", 10)
add_normal("▸ BI harus pertahankan atau naikkan bunga untuk jaga inflasi — pertumbuhan ekonomi tertekan.", 10)
add_normal("▸ Kondisi ini menciptakan dilema: tahan inflasi (bunga naik) atau growth first (inflasi meledak).", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Inflasi global bertahan = Indonesia tidak bisa turunkan bunga terlalu cepat.", 10)
add_normal("● \"Double scar\" bukan hanya masalah Barat — Indonesia juga merasakan dampak berganda.", 10)
add_normal("● Kebijakan anti-inflasi yang terlalu ketat akan menekan pertumbuhan yang sudah lemah.", 10)
add_normal("● Diperlukan kebijakan yang lebih smart dan targeted — bukan blanket restriction.", 10)
doc.add_paragraph()

# BERITA 4
add_normal("Berita 4   Pasar Modal · Global", 11, True)
add_normal("Tech Rally Global — Nasdaq Raup Gain 8% Mei, Saham Software Best Month Since 2001", 10, True)
add_normal("25–30 Mei 2026    |  Pasar Modal  Global", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Nasdaq mencatat gain 8% di bulan Mei — ditutup di level tertinggi sepanjang masa.", 10)
add_normal("▸ Saham software mencatat bulan terbaik sejak 2001 — \"SaaSpocalypse\" mereda.", 10)
add_normal("▸ Dell melonjak 32% dalam sehari — penjualan AI server yang booming.", 10)
add_normal("▸ Wall Street rally ke Juni dengan harapan kesepakatan AS-Iran.", 10)
add_normal("▸ Namun rally ini tidak berdampak signifikan pada pasar emerging markets seperti Indonesia.", 10)
add_normal("▸ Capital tetap mengalir keluar Indonesia meskipun pasar AS rally — sinyal negative untuk IHSG.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Rally global tidak otomatis mengalir ke Indonesia — kita perlu perbaiki fundamental.", 10)
add_normal("● Nasdaq rally tapi IHSG tetap melemah = masalah struktural pasar modal kita.", 10)
add_normal("● Teknologi dan AI adalah masa depan — Indonesia perlu siapkan ekosistem yang mendukung.", 10)
add_normal("● Tanpa reformasi, kita akan terus tertinggal dari rally global.", 10)
doc.add_paragraph()

# BERITA 5
add_normal("Berita 5   Pasar · Global", 11, True)
add_normal("Kospi & Topix Rekor Baru — Investor Abaikan Ketegangan Iran, Emerging Markets Mixed", 10, True)
add_normal("25–30 Mei 2026    |  Pasar  Global", 9)
doc.add_paragraph()
add_normal("FAKTA UTAMA", 10, True)
add_normal("▸ Korea Kospi dan Japan Topix mencapai level tertinggi baru.", 10)
add_normal("▸ Investor global mengabaikan ketegangan geopolitik Iran — fokus pada earnings.", 10)
add_normal("▸ Kondisi emerging markets mixed — beberapa negara rally, lainnya tertekan.", 10)
add_normal("▸ Indonesia termasuk yang tertekan — capital outflow berlanjut.", 10)
add_normal("▸ Korea dan Japan menunjukkan bahwa emerging markets bisa rally jika fundamental kuat.", 10)
add_normal("▸ Indonesia perlu belajar dari keberhasilan negara tetangga.", 10)
doc.add_paragraph()
add_normal("HOLDING STATEMENT FRAKSI PENYEIMBANG", 10, True)
add_normal("● Kospi dan Topix rekor = bukti emerging markets bisa sukses dengan fundamental benar.", 10)
add_normal("● Indonesia tertinggal bukan karena kondisi global — tapi karena masalah domestik.", 10)
add_normal("● Reformasi struktural diperlukan — bukan hanya kebijakan taktis.", 10)
add_normal("● Kita perlu peta jalan yang jelas untuk pulihkan kepercayaan investor.", 10)
doc.add_page_break()

# ==================== PAGE 3 ====================
add_title("Screening Berita Ekonomi Global — François Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("François Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 3 dari 4", 9)
doc.add_paragraph()

add_title("SUMBER & CATATAN KAKI", 12)

add_normal("¹ CNBC, 25–30 Mei 2026 — \"Brent crude biggest monthly loss since 2020\"", 9)
add_normal("² CNBC, 25–30 Mei 2026 — \"US household paying $450 more on gas\"", 9)
add_normal("³ CNBC, 25–30 Mei 2026 — \"Fed holds rates, market expectations\"", 9)
add_normal("⁴ CNBC, 25–30 Mei 2026 — \"Dollar remains strong\"", 9)
add_normal("⁵ CNBC, 25–30 Mei 2026 — \"UK household energy bills jump 13%\"", 9)
add_normal("⁶ CNBC, 25–30 Mei 2026 — \"Double scar inflation\"", 9)
add_normal("⁷ CNBC, 25–30 Mei 2026 — \"Nasdaq gained 8% in May\"", 9)
add_normal("⁸ CNBC, 25–30 Mei 2026 — \"Software stocks best month since 2001\"", 9)
add_normal("⁹ CNBC, 25–30 Mei 2026 — \"Dell stock skyrocketed 32%\"", 9)
add_normal("¹⁰ CNBC, 25–30 Mei 2026 — \"Kospi and Topix reached new peaks\"", 9)

doc.add_page_break()

# ==================== PAGE 4 ====================
add_title("Screening Berita Ekonomi Global — François Penyeimbang DPR RI   ·   25–30 Mei 2026", 10)
add_centered("François Penyeimbang DPR RI  ·  Dokumen Riset Internal  ·  Halaman 4 dari 4", 9)
doc.add_paragraph()

add_title("RINGKASAN EKSEKUTIF", 12)

add_normal("Periode pelaporan: 25–30 Mei 2026", 10)
add_normal("Topik dominan: Oil Prices · Fed Rates · Global Inflation · Trade · Markets", 10)
add_normal("Skala tekanan: 🔴 Sangat Tinggi", 10)
doc.add_paragraph()

add_normal("Lima berita ekonomi global terpanas periode ini:", 10, True)
add_normal("1. Oil Shock — Brent loss terbesar since 2020, household AS extra $450", 10)
add_normal("2. Fed Holds Rates — Dollar kuat tekan rupiah", 10)
add_normal("3. Inflasi Global — UK energy bills +13%, \"double scar\"", 10)
add_normal("4. Tech Rally — Nasdaq +8%, software best month since 2001", 10)
add_normal("5. Kospi & Topix Rekor — investor abaikan Iran", 10)
doc.add_paragraph()

add_normal("Dampak terhadap Indonesia:", 10, True)
add_normal("• Oil shock → ICP jebol → beban subsidi energi membengkak", 10)
add_normal("• Fed hawkish → rupiah lemah → beban utang valas naik", 10)
add_normal("• Inflasi global → BI harus naikkan bunga → pertumbuhan tertekan", 10)
add_normal("• Rally global tidak mengalir ke Indonesia → capital outflow", 10)
add_normal("• Kospi/Topix rekor menunjukkan emerging markets bisa sukses", 10)
doc.add_paragraph()

add_normal("Kesimpulan:", 10, True)
add_normal("Ekonomi global memberikan tekanan multi-dimensi pada fiskal dan moneter Indonesia. Diperlukan kebijakan yang lebih proaktif dan reformasi struktural — bukan hanya respons taktis. Pelajaran dari negara tetangga (Korea, Japan) harus ditiru.", 10)

# Save
output_path = r"C:\Users\USER\Screening_Ekonomi_Global_25-30Mei2026.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")