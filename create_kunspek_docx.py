from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper functions
def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_photo(path, width=Inches(4)):
    try:
        doc.add_picture(path, width=width)
    except Exception as e:
        print(f"Error adding photo: {e}")

# ==================== PAGE 1: SAMPUL ====================
add_title("LAPORAN KUNJUNGAN SPESIFIK", 20)
add_title("KE - 2", 20)
add_title("TAHUN SIDANG 2023-2024", 16)

doc.add_paragraph()
doc.add_paragraph()

add_centered("Nama                        : Ir. Andreas Eddy Susetyo, M.M.")
add_centered("Nomor Anggota      : A-214")
add_centered("Komisi                       : XI")
add_centered("Fraksi                        : PDI-Perjuangan")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 21 Februari - 23 Februari 2024")

doc.add_page_break()

# ==================== PAGE 2: PENDAHULUAN ====================
add_title("1. PENDAHULUAN", 14)

pendahuluan_text = """Kunjungan kerja Anggota DPR RI untuk Dapil Jawa Timur V Malang Raya mempunyai beberapa fungsi, antara lain, pertama, sebagai media penyerapan aspirasi masyarakat paling bawah yang berada di Kota Malang, Kota Batu dan Kabupaten Malang. Kunjungan kerja ini akan dijadikan masukan bagi Fraksi PDI-Perjuangan terkait dengan usulan dan keinginan masyarakat bawah.

Kedua, dengan kunjungan ini ternyata sangat diharapkan oleh masyarakat bawah yang merasa mempunyai wakil yang duduk di Lembaga Legislatif. Kehadiran Anggota DPR RI ini dapat mempererat hubungan antara konstituen dengan Anggota DPR.

Ketiga, dengan kunjungan ini diharapkan DPR lebih sensitif dengan permasalahan yang dihadapi oleh masyarakat bawah dan utamanya segala permasalahan yang terjadi di tubuh struktural PDI-Perjuangan yang ada di Malang Raya.

Berbeda dengan kunjungan kerja reses pada umumnya yang bersifat umum dan menjangkau berbagai isu sekaligus, Kunjungan Spesifik dirancang sebagai bentuk pendalaman tematik terhadap satu sektor atau isu tertentu yang dianggap strategis dan memerlukan perhatian khusus dari Anggota DPR RI. Pendekatan tematik ini memungkinkan penggalian informasi yang lebih mendalam, komprehensif, dan berbasis data langsung dari lapangan, sehingga rekomendasi kebijakan yang dihasilkan dapat lebih tepat sasaran dan aplikatif bagi penyelesaian permasalahan yang dihadapi oleh kelompok masyarakat terkait.

Pada kesempatan kali ini, Kunjungan Spesifik Ke-2 Tahun Sidang 2023-2024 difokuskan pada sektor Usaha Mikro, Kecil, dan Menengah (UMKM), mengingat sektor ini memiliki kontribusi yang sangat signifikan terhadap perekonomian nasional maupun regional, khususnya dalam hal penciptaan lapangan kerja, pemerataan pendapatan, dan penguatan ketahanan ekonomi masyarakat di tingkat akar rumput. Sebagai Anggota Komisi XI DPR RI yang membidangi urusan keuangan, perbankan, dan perencanaan pembangunan, Ir. Andreas Eddy Susetyo, M.M. memandang penting untuk mendalami secara langsung kondisi riil pelaku UMKM di Malang Raya, termasuk tantangan permodalan, akses pasar, dan kesiapan digitalisasi usaha.

Malang Raya dikenal sebagai salah satu kawasan dengan ekosistem UMKM yang cukup dinamis, mulai dari kerajinan batik, produk olahan pangan, hingga usaha kreatif berbasis pariwisata. Keberagaman sektor UMKM ini menjadikan wilayah tersebut representatif untuk dijadikan lokasi kajian mendalam mengenai efektivitas kebijakan pemberdayaan UMKM yang telah dan sedang berjalan, sekaligus mengidentifikasi kebutuhan riil pelaku usaha yang belum sepenuhnya terakomodasi oleh program pemerintah yang ada saat ini.

Berbagai tantangan yang umum dihadapi pelaku UMKM, seperti keterbatasan akses pembiayaan yang terjangkau, minimnya pemahaman terhadap strategi pemasaran digital, serta kurangnya pendampingan teknis dalam pengembangan produk, menjadi fokus utama penggalian informasi dalam kunjungan ini. Dialog langsung dengan para pelaku UMKM, kelompok perempuan pengusaha, dan komunitas ekonomi kreatif diharapkan dapat menghasilkan pemetaan permasalahan yang akurat sebagai dasar perumusan rekomendasi kebijakan yang lebih responsif terhadap kebutuhan pelaku usaha kecil di daerah.

Dengan demikian, laporan Kunjungan Spesifik ini disusun tidak hanya sebagai bentuk dokumentasi kegiatan, tetapi juga sebagai instrumen pertanggungjawaban dan bahan evaluasi bagi perumusan kebijakan pemberdayaan UMKM yang lebih tepat sasaran di tingkat nasional maupun daerah. Melalui pendekatan tematik ini, diharapkan setiap rekomendasi yang dihasilkan benar-benar berpijak pada data dan aspirasi riil dari pelaku usaha di lapangan, bukan sekadar asumsi kebijakan yang dirumuskan dari tingkat pusat tanpa verifikasi langsung terhadap kondisi aktual masyarakat."""

for paragraph in pendahuluan_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
add_centered("Malang, 24 Februari 2024")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("Nomor Anggota A-214")

doc.add_page_break()

# ==================== RANGKAIAN KEGIATAN ====================
add_title("2. RANGKAIAN KEGIATAN", 14)

photo_folder = r"D:\DATA PAPA\LAPORAN\KUNSPEK\KUNSPEK 2"

# KEGIATAN 1
add_title("KEGIATAN 1", 12)
keg1_text = """Kunjungan Kerja Spesifik kali ini terkait Usaha Mikro, Kecil, dan Menengah (UMKM) dimana Andreas menjelaskan bahwa UMKM memiliki kontribusi yang signifikan dalam perekonomian, terutama dalam menciptakan lapangan kerja, mempromosikan inovasi, dan meningkatkan pertumbuhan ekonomi di tingkat lokal dan nasional. UMKM mencakup berbagai sektor ekonomi, mulai dari perdagangan, industri, hingga jasa, dan menjadi tulang punggung ekonomi di banyak negara, termasuk Indonesia.

UMKM berperan penting dalam menciptakan lapangan kerja. Sektor UMKM memberikan kesempatan kerja kepada banyak orang, terutama di wilayah pedesaan dan perkotaan yang mungkin memiliki akses terbatas ke pekerjaan formal. Melalui kegiatan produksi, distribusi, dan pemasaran, UMKM memberikan pekerjaan kepada banyak individu, membantu mengurangi tingkat pengangguran, dan meningkatkan taraf hidup masyarakat."""
add_normal(keg1_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K1.JPG"), width=Inches(5.5))
p = add_centered("Andreas Eddy Susetyo bersama pengrajin batik di Jl. Letjend S. Parman No.44, Purwantoro, Kec. Blimbing, Kota Malang, Jawa Timur", size=9)

doc.add_page_break()

# KEGIATAN 2
add_title("KEGIATAN KE - 2", 12)
keg2_text = """Selain itu, UMKM juga dapat menjadi agen inklusivitas ekonomi. Banyak UMKM dimiliki oleh kelompok masyarakat yang sebelumnya mungkin tidak memiliki akses ke sektor formal, seperti perempuan, pemuda, dan kelompok minoritas. Dengan memberikan peluang usaha kepada kelompok-kelompok ini, UMKM berkontribusi pada pemberdayaan ekonomi dan sosial, menciptakan inklusivitas yang diperlukan untuk pembangunan berkelanjutan.

Andreas Eddy Susetyo anggota DPR/MPR-RI Fraksi PDI-Perjuangan dari Dapil Jawa Timur V menegaskan bahwa UMKM juga memiliki dampak positif terhadap pemerataan ekonomi. Dengan tersebarnya UMKM di berbagai daerah, terutama di wilayah yang sebelumnya mungkin tidak mendapatkan manfaat dari pembangunan ekonomi, UMKM dapat membantu mengurangi kesenjangan ekonomi antar wilayah.

Untuk meningkatkan kontribusi UMKM dalam menciptakan lapangan kerja, mempromosikan inovasi, dan meningkatkan pertumbuhan ekonomi, penting bagi pemerintah dan pemangku kepentingan lainnya untuk memberikan dukungan yang memadai. Ini termasuk penyediaan pelatihan, pembiayaan yang terjangkau, akses ke pasar, dan kebijakan yang mendukung perkembangan UMKM."""
add_normal(keg2_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K2.JPG"), width=Inches(5.5))
add_centered("Andreas dalam kunjungannya dalam acara pameran produk UMKM di Jl. Besar Ijen, Oro-oro Dowo, Kec. Klojen, Kota Malang, Jawa Timur", size=9)
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K3.JPG"), width=Inches(5.5))

doc.add_page_break()

# ==================== PENUTUP ====================
add_title("3. PENUTUP", 14)

penutup_text = """Dalam kunjungan kerja Anggota DPR RI untuk Dapil Jawa Timur V Malang Raya, terungkap dengan jelas bahwa UMKM memiliki peran krusial dalam menggerakkan roda perekonomian di tingkat lokal dan nasional. Andreas Eddy Susetyo dengan tegas menggarisbawahi kontribusi UMKM dalam menciptakan lapangan kerja, mempromosikan inovasi, dan meningkatkan pertumbuhan ekonomi. Kunjungan ini tidak hanya memenuhi fungsi sebagai medium penyerapan aspirasi masyarakat paling bawah, tetapi juga menyoroti pentingnya dukungan pemerintah dan pemangku kepentingan lainnya untuk memastikan kesinambungan dan pengembangan UMKM. Secara keseluruhan, UMKM menjadi tulang punggung bagi pembangunan ekonomi yang inklusif dan berkelanjutan.

Dari hasil dialog langsung dengan para pelaku UMKM di sentra kerajinan batik maupun peserta pameran produk UMKM di Malang Raya, diperoleh gambaran yang lebih utuh mengenai kondisi riil di lapangan. Sebagian besar pelaku usaha menyampaikan bahwa akses terhadap pembiayaan yang terjangkau masih menjadi kendala utama dalam mengembangkan skala usaha mereka. Selain itu, keterbatasan pengetahuan mengenai strategi pemasaran digital turut menghambat perluasan jangkauan pasar produk UMKM, terutama bagi pelaku usaha yang berasal dari kelompok perempuan dan generasi muda yang baru merintis usaha.

Sebagai Anggota Komisi XI DPR RI, temuan-temuan dari Kunjungan Spesifik ini akan menjadi bahan masukan penting dalam pembahasan kebijakan pemberdayaan UMKM, khususnya terkait skema pembiayaan mikro yang lebih inklusif, penguatan program pelatihan kewirausahaan, serta pengembangan infrastruktur digital yang mendukung transformasi UMKM menuju ekonomi digital. Komitmen untuk terus mengawal implementasi program-program pemberdayaan UMKM di tingkat nasional akan dijalankan secara berkesinambungan, sejalan dengan aspirasi yang telah disampaikan langsung oleh para pelaku usaha di Malang Raya.

Keberadaan UMKM yang tersebar di berbagai wilayah, mulai dari sentra kerajinan batik di Kecamatan Blimbing hingga pelaku usaha kreatif di kawasan Kecamatan Klojen, menegaskan bahwa potensi ekonomi kerakyatan di Malang Raya sangat besar dan perlu terus didorong melalui kebijakan yang berpihak. Sinergi antara pemerintah pusat, pemerintah daerah, perbankan, dan lembaga pendamping UMKM menjadi kunci utama agar potensi tersebut dapat dikembangkan secara optimal dan berkelanjutan, sehingga mampu meningkatkan kesejahteraan pelaku usaha kecil secara nyata.

Kunjungan Spesifik ini juga menegaskan pentingnya pendekatan tematik dalam menjalankan fungsi pengawasan dan aspirasi Anggota DPR RI. Dengan memfokuskan perhatian pada satu sektor secara mendalam, permasalahan yang selama ini mungkin luput dari perhatian dalam kunjungan kerja reses yang bersifat umum dapat teridentifikasi secara lebih spesifik dan komprehensif. Pendekatan semacam ini diharapkan dapat terus dikembangkan dan diterapkan pada sektor-sektor strategis lainnya di Dapil Jawa Timur V, seperti pertanian, pariwisata, dan ekonomi kreatif, guna menghasilkan rekomendasi kebijakan yang lebih berbasis data dan kebutuhan riil masyarakat.

Akhir kata, laporan Kunjungan Spesifik Ke-2 ini diharapkan dapat menjadi rujukan yang bermanfaat bagi perumusan kebijakan pemberdayaan UMKM yang lebih tepat sasaran, sekaligus menjadi wujud komitmen nyata Anggota DPR RI dalam menjalankan fungsi pengawasan dan aspirasi masyarakat di Dapil Jawa Timur V secara berkelanjutan. Semoga sinergi yang telah terjalin antara Anggota Dewan, pemerintah daerah, dan pelaku UMKM dapat terus terjaga demi kemajuan ekonomi kerakyatan yang inklusif dan berkeadilan."""

for paragraph in penutup_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 24 Februari 2024")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("Nomor Anggota A-214")

# Save document
output_path = r"D:\DATA PAPA\LAPORAN\KUNSPEK\KUNSPEK 2\Laporan Kunjungan Spesifik 2 2024-1.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
