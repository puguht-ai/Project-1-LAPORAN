from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper functions
def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_photo(path, width=Inches(4)):
    try:
        doc.add_picture(path, width=width)
    except Exception as e:
        print(f"Error adding photo: {e}")

# ==================== PAGE 1: SAMPUL ====================
add_title("LAPORAN KUNJUNGAN KERJA", 20)
add_title("DAERAH PILIHAN", 20)
add_title("KE - VII", 20)
add_title("TAHUN SIDANG 2025-2026", 20)

doc.add_paragraph()
doc.add_paragraph()

add_centered("Nama                        : Ir. Andreas Eddy Susetyo, M.M.")
add_centered("Nomor Anggota      : A-220")
add_centered("Komisi                       : XI")
add_centered("Fraksi                        : PDI-Perjalanan")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 29 – 31 Mei 2026")

doc.add_page_break()

# ==================== PAGE 2: PENDAHULUAN ====================
add_title("1. PENDAHULUAN", 14)

pendahuluan_text = """Sektor pariwisata memiliki potensi yang sangat besar untuk menjadi pendorong utama pertumbuhan ekonomi masyarakat di wilayah Malang Raya. Sektor ini tidak hanya sekadar menjadi magnet bagi wisatawan, tetapi di dalamnya terdapat ekosistem perekonomian yang sangat potensial dan saling terkait erat dengan berbagai sektor pendukung lainnya. Ketika wisatawan mengunjungi suatu daerah, mereka tidak hanya mengeluarkan biaya untuk门票 atau tarif masuk objek wisata, tetapi juga mengonsumsi berbagai produk dan jasa pendukung yang menggerakkan roda perekonomian secara menyeluruh. Hal ini mencakup sektor UMKM makanan dan minuman yang menyediakan kuliner khas daerah, jasa transportasi baik darat maupun udara yang mengangkut wisatawan dari dan menuju destinations, ekonomi kreatif seperti oleh-oleh khas daerah, kerajinan tangan, dan berbagai produk unik yang menjadi identitas budaya lokal, serta sektor pendukung lainnya seperti penginapan, restoran, dan layanan guides profesional yang memberikan pengalaman tak terlupakan bagi para pengunjung. Dengan demikian, pengembangan sektor pariwisata secara otomatis akan memberikan dampak berganda (multiplier effect) yang positif bagi pertumbuhan ekonomi masyarakat di berbagai lapisan.

Berdasarkan data Badan Pusat Statistik (BPS), kontribusi sektor pariwisata terhadap Produk Domestik Bruto (PDB) nasional pada tahun 2025 tercatat sebesar 4,12%, mengalami peningkatan dari 3,86% pada tahun sebelumnya. Sektor ini menyerap tenaga kerja langsung sebesar 5,2 juta orang, atau sekitar 4,3% dari total angkatan kerja nasional. Kontribusi Devisa dari sektor pariwisata pada tahun 2025 mencapai USD 14,8 miliar, meningkat 12,4% dari tahun sebelumnya. Angka ini menunjukkan bahwa sektor pariwisata memiliki peran strategis dalam neraca perdagangan nasional dan menjadi salah satu sektor andalan dalam menghasilkan devisa negara. Pertumbuhan sektor pariwisata nasional yang positif ini menjadi modal penting dalam membangun kepercayaan investor dan mendorong pengembangan infrastruktur pendukung pariwisata di berbagai daerah.

Di tingkat regional, Malang Raya yang mencakup Kota Malang, Kabupaten Malang, dan Kota Batu memiliki potensi pariwisata yang sangat luar biasa dan belum sepenuhnya dimaksimalkan. Data Dinas Pariwisatan Jawa Timur mencatat bahwa kontribusi sektor pariwisata terhadap PDRB Kabupaten Malang pada tahun 2025 mencapai 8,7%, sementara Kota Malang mencatat 6,3% dan Kota Batu mencapai 15,2%. Kota Batu sebagai kota wisata memiliki kontribusi tertinggi karena hampir seluruh wilayahnya dikembangkan sebagai destinations pariwisata. Secara keseluruhan, kontribusi sektor pariwisata di Malang Raya terhadap PDRB regional mencapai rata-rata 8,4%, jauh di atas rata-rata nasional sebesar 4,12%. Hal ini menunjukkan bahwa Malang Raya memiliki keunggulan komparatif yang sangat signifikan di sektor pariwisata dibandingkan wilayah lain di Indonesia. Potensi ini harus dimaksimalkan melalui berbagai strategi pengembangan yang komprehensif dan berkelanjutan.

Namun, di balik potensi yang besar tersebut, sektor pariwisata di Malang Raya masih menghadapi berbagai tantangan yang perlu segera diatasi. Beberapa problematika utama yang dihadapi antara lain adalah keterbatasan infrastruktur pendukung seperti jalan akses menuju beberapa destinations wisata yang masih rusak dan kurang memadai, minimnya promosi dan pemasaran digital yang membuat banyak potensi wisata tersembunyi dari mata dunia, kurangnya kualitas sumber daya manusia di sektor hospitality dan layanan pariwisata, serta kurangnya koordinasi antara pemerintah daerah, pelaku usaha, dan masyarakat dalam pengembangan sektor pariwisata secara terpadu. Kunjungan Daerah Pemilihan (Kundapil) kali ini dirancang untuk mendengarkan langsung berbagai masukan dari para stakeholder pariwisata, mengevaluasi kondisi riil di lapangan, serta merumuskan strategi penanganan yang konkret. Dengan melibatkan berbagai pihak seperti Dinas Pariwisatan, pelaku usaha pariwisata, asosiasi hotel dan restoran, serta masyarakat lokal, diharapkan dapat ditemukan solusi yang tepat untuk mengembangkan potensi pariwisata di Malang Raya secara maksimal."""

for paragraph in pendahuluan_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_page_break()

# ==================== RANGKAIAN KEGIATAN ====================
add_title("2. RANGKAIAN KEGIATAN", 14)

photo_folder = r"D:\DATA PAPA\LAPORAN\KUNDAPIL\KUNDAPIL 7"

# KEGIATAN 1 - Full Page
add_title("KEGIATAN 1", 12)
keg1_text = """Menghadiri presentasi dari pelaku pariwisata Pantai Sendang Biru dan sekitarnya di Sendang Biru di Desa Sumberagung, Kec. Sumbermanjing Wetan, Kab. Malang. Pantai Sendang Biru merupakan salah satu destinations andalan Malang Raya yang memiliki keindahan alam laut yang memukau dengan hamparan pasir putih dan air laut yang jernih. Para pelaku pariwisata menyampaikan berbagai masukan mengenai perlunya peningkatan infrastruktur akses jalan menuju pantai, penyediaan fasilitasMCK yang lebih memadai, serta peningkatan keamanan bagi wisatawan. Mereka juga menyampaikan harapan agar pemerintah dapat membantu pemasaran digital agar semakin banyak wisatawan yang mengetahui keberadaan destinations ini."""
add_normal(keg1_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K1.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K2.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 2 - Full Page
add_title("KEGIATAN 2", 12)
keg2_text = """Berdialog dengan Dinas Pariwisatan Kota Malang untuk mencari masukan tentang berbagai potensi dan problem yang dihadapi sektor pariwsata Kota Malang. Dalam pertemuan yang berlangsung produktif ini, Dinas Pariwisatan menyampaikan berbagai rencana pengembangan destinations pariwisata baru, problematikapenyediaan air bersih di beberapa destinations wisata, serta kebutuhan akan pelatihan sumber daya manusia di sektor hospitality. Mereka menekankan perlunya sinergi antara pemerintah kota dengan kabupaten tetangga untuk mengembangkan Malang Raya sebagai satu kawasan wisata yang terintegrasi."""
add_normal(keg2_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K3.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K4.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 3 - Full Page
add_title("KEGIATAN 3", 12)
keg3_text = """Kunjungan Lapangan ke daerah tujuan wisata Pantai Asmara, Desa Tambakjo, Kecamatan Sumbermanjing Wetan, Kab. Malang. Pantai Asmara yang dikenal dengan sebutan "Raja Ampat versi Malang" ini memiliki potensi keindahan bawah laut yang sangat luar biasa. Dalam kunjungan ini, rombongan melihat langsung kondisi infrastruktur pendukung dan berdiskusi dengan masyarakat lokal mengenai pengembangan ekowisata bahari yang berkelanjutan. Para pelaku tourism menyampaikan kebutuhan akan bantuan peralatan selam dan snorkeling untuk mengembangkan aktivitas diving dan snorkeling yang aman bagi wisatawan."""
add_normal(keg3_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K5.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K6.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 4 - Full Page
add_title("KEGIATAN 4", 12)
keg4_text = """Bertemu para pelakuUMKM sektor pariwsata dari wilayah Kota Batu dan Kab. Malang. Pertemuan ini menjadi forum strategis untuk membahas berbagai upaya penguatan ekonomi kreatif di sektor pariwsata. Para pelakuUMKM menyampaikan keluhan mengenai minimnya akses pemasaran, keterbatasan modal untuk mengembangkan produksi, serta kebutuhan akan pelatihan strategi pemasaran digital. Mereka menekankan perlunya platform pemasaran bersama yang dapat mempertemukan pelakuUMKM dengan wisatawan secara langsung."""
add_normal(keg4_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K7.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K8.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 5 - Full Page
add_title("KEGIATAN 5", 12)
keg5_text = """Menyerahkan penghargaan dari Bank Indonesia kepada pelaku pariwsata Malang Raya yang telah memanfaatkan fasilitas pembayaran digital QRIS. Penghargaan ini merupakan apresiasi atas upaya digitalisasi pembayaran yang dilakukan oleh pelaku pariwsata dalam rangka mendukung Gerakan Nasional Non Tunai (GNNT). Para penerima penghargaan menyampaikan komitmen mereka untuk terus meningkatkan kualitas layanan dan memperluas penggunaan pembayaran digital."""
add_normal(keg5_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K9.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K10.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 6 - Full Page
add_title("KEGIATAN 6", 12)
keg6_text = """Berdialog dengan pengelola obyek wisata Coban Putri di Desa Krajan, Dusun Tlekung (Oro-Oro Ombo), Kecamatan Junrejo, Kota Batu. Coban Putri merupakan salah satu destinations andalan Kota Batu yang memiliki keindahan air terjun di tengah hutan pinus. Dalam dialog ini, pengelola menyampaikan berbagai tantangan dalam pengelolaan destinations, kebutuhan akan perbaikan infrastruktur parkiran, serta rencana pengembangan wahana baru yang lebih menarik bagi wisatawan. Mereka juga menyampaikan harapan agar pemerintah dapat membantu promosi melalui berbagai media."""
add_normal(keg6_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K11.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K12.JPG"), width=Inches(5.5))
doc.add_page_break()

# KEGIATAN 7 - Full Page
add_title("KEGIATAN 7", 12)
keg7_text = """Menyaksikan Pergelaran Seni Budaya Ludruk Kenduri Lawasan di Kota Malang. Banyak seni pertunjukan di Malang Raya yang belum dimaksimalkan menjadi asset wisata potensial. Pertunjukan ludruk ini menampilkan berbagai lakon tradisional yang menggambarkan kearifan lokal masyarakat Jawa Timur. Dalam kesempatan ini, dibicarakan perlunya pengembangan seni pertunjukan tradisional sebagai daya tarik wisata budaya yang dapat menarik wisatawan domestik maupun mancanegara. Para seniman menyampaikan kebutuhan akan dukungan pemerintah dalam bentuk fasilitas latihan dan panggung pertunjukan yang memadai."""
add_normal(keg7_text)

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K13.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "K14.JPG"), width=Inches(5.5))
doc.add_page_break()

# ==================== PENUTUP ====================
add_title("3. PENUTUP", 14)

penutup_text = """Rangkaian Kunjungan Daerah Pemilihan yang telah berlangsung selama tiga hari di Malang Raya telah memberikan pemahaman mendalam mengenai potensi dan tantangan yang dihadapi sektor pariwisata di wilayah ini. Data yang diperoleh menunjukkan bahwa sektor pariwisata memiliki kontribusi yang signifikan terhadap PDRB regional, yaitu rata-rata 8,4%, jauh di atas rata-rata nasional sebesar 4,12%. Malang Raya memiliki keunggulan komparatif yang sangat besar dengan berbagai destinations wisata alam, budaya, dan buatan yang menarik. Namun, potensi tersebut masih belum dimaksimalkan secara optimal karena berbagai kendala yang dihadapi oleh para stakeholder pariwisata.

Kunjungan ke berbagai lokasi menegaskan bahwa infrastruktur pendukung menjadi salah satu aspek kritis yang perlu segera diperbaiki. Jalan akses menuju beberapa destinations wisata, terutama di kawasan pantai seperti Pantai Sendang Biru dan Pantai Asmara, masih memerlukan perbaikan signifikan. Selain itu, penyediaan fasilitasMCK yang memadai, peningkatan kualitas air bersih, dan perbaikan infrastruktur listrik menjadi prioritas utama yang perlu mendapat perhatian dari pemerintah daerah. Tanpa infrastruktur yang memadai, potensi pariwisata yang besar tidak akan dapat dikembangkan secara maksimal dan akan mengecewakan wisatawan yang mengunjungi destinations tersebut.

Pengembangan ekonomi kreatif di sektor pariwisata juga menjadi temuan penting dari rangkaian kunjungan ini. Para pelakuUMKM di sektor pariwisata memiliki potensi yang sangat besar untuk meningkatkan kesejahteraan mereka melalui pengembangan produk-produk kreatif yang unik dan berbudaya lokal. Namun, mereka menghadapi berbagai tantangan seperti minimnya akses pemasaran, keterbatasan modal, dan kurangnya pelatihan dalam strategi pemasaran digital. Pendampingan dan pelatihan intensif menjadi kebutuhan mendesak yang harus dipenuhi untuk memberdayakan para pelakuUMKM ini. Dengan pengembangan ekonomi kreatif yang tepat, sektor pariwisata dapat menjadi mesin penggerak utama pertumbuhan ekonomi masyarakat di Malang Raya.

Kolaborasi antara pemerintah, pelaku usaha, dan masyarakat menjadi fondasi utama dalam membangun ketahanan ekonomi di sektor pariwisata. Sinergi antara berbagai stakeholder diperlukan untuk mengembangkan potensi pariwisata secara terpadu dan berkelanjutan. Pemerintah daerah perlu meningkatkan koordinasi dalam pengembangan infrastruktur dan promosi, pelaku usaha perlu meningkatkan kualitas layanan dan inovasi produk, sementara masyarakat perlu dilibatkan secara aktif dalam pengembangan destinations pariwisata sebagai pemilik dan penjaga kelestarian budaya lokal. Dengan kolaborasi yang solid, Malang Raya dapat menjadi destinations pariwisata kelas dunia yang menarik wisatawan domestik maupun mancanegara.

Komitmen untuk terus menyerap aspirasi dan memperjuangkan kebijakan yang berpihak pada masyarakat menjadi prioritas utama dalam setiap kunjungan kerja. Setiap masukan yang diterima akan menjadi bahan penting dalam perumusan kebijakan di tingkat nasional. Dengan sinergi antara pemerintah, pelaku usaha, dan masyarakat, diharapkan Malang Raya dapat segera pulih dan kembali menjadi pusat pariwisata yang tangguh dan berkelanjutan. Sektor pariwisata memiliki potensi yang sangat besar untuk menjadi motor penggerak utama ekonomi masyarakat, dan potensi ini harus dimaksimalkan melalui berbagai kebijakan yang mendukung dan pemberdayaan yang berkelanjutan."""

for paragraph in penutup_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 1 Juni 2026")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("No Anggota: A-220")

# Save document
output_path = r"D:\DATA PAPA\LAPORAN\KUNDAPIL\KUNDAPIL 7\KUNDAPIL VII 29_31 MEI 2026_v2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")