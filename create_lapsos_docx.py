from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# Helper functions
def add_title(text, size=16, bold=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_normal(text, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_centered(text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_photo(path, width=Inches(4)):
    try:
        doc.add_picture(path, width=width)
    except Exception as e:
        print(f"Error adding photo: {e}")

# ==================== PAGE 1: SAMPUL ====================
add_title("LAPORAN SOSIALISASI UNDANG-UNDANG", 20)
add_title("KE - 1", 20)

doc.add_paragraph()
doc.add_paragraph()

add_centered("Nama                        : Ir. Andreas Eddy Susetyo, M.M.")
add_centered("Nomor Anggota      : A-220")
add_centered("Komisi                       : XI")
add_centered("Fraksi                        : PDI-Perjuangan")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 09 - 10 April 2025")

doc.add_page_break()

# ==================== PAGE 2: PENDAHULUAN ====================
add_title("1. PENDAHULUAN", 14)

pendahuluan_text = """Kunjungan kerja Anggota DPR RI untuk Dapil Jawa Timur V Malang Raya mempunyai beberapa fungsi, antara lain, pertama, sebagai media penyerapan aspirasi masyarakat paling bawah yang berada di Kota Malang, Kota Batu dan Kabupaten Malang. Kunjungan kerja ini akan dijadikan masukan bagi Fraksi PDI-Perjuangan terkait dengan usulan dan keinginan masyarakat bawah.

Kedua, dengan kunjungan ini ternyata sangat diharapkan oleh masyarakat bawah yang merasa mempunyai wakil yang duduk di Lembaga Legislatif. Kehadiran Anggota DPR RI ini dapat mempererat hubungan antara konstituen dengan Anggota DPR.

Ketiga, dengan kunjungan ini diharapkan DPR lebih sensitif dengan permasalahan yang dihadapi oleh masyarakat bawah dan utamanya segala permasalahan yang terjadi di tubuh struktural PDI-Perjuangan yang ada di Malang Raya.

Selain fungsi penyerapan aspirasi, kunjungan kerja ini juga menjalankan salah satu fungsi konstitusional DPR RI yaitu fungsi legislasi, khususnya dalam hal sosialisasi produk perundang-undangan kepada masyarakat di daerah pemilihan. Sosialisasi Undang-Undang merupakan tahap krusial dalam siklus kebijakan publik, karena sebaik apapun sebuah regulasi dirumuskan, manfaatnya tidak akan optimal apabila masyarakat, pelaku usaha, dan aparatur di tingkat daerah tidak memahami substansi, tujuan, dan mekanisme pelaksanaannya. Oleh karena itu, kegiatan sosialisasi ini dirancang sebagai jembatan komunikasi dua arah antara pembuat kebijakan di tingkat pusat dengan pemangku kepentingan di lapangan, sehingga setiap ketentuan yang berlaku dapat dipahami secara utuh dan diimplementasikan sesuai dengan semangat yang terkandung di dalamnya.

Sebagai Anggota Komisi XI DPR RI yang membidangi urusan keuangan, perbankan, dan perencanaan pembangunan nasional, serta menjabat sebagai Ketua Badan Akuntabilitas Keuangan Negara (BAKN), Ir. Andreas Eddy Susetyo, M.M. memiliki tanggung jawab khusus untuk memastikan bahwa setiap regulasi terkait pengelolaan keuangan negara dan program strategis pemerintah dapat diimplementasikan secara transparan, akuntabel, dan tepat sasaran hingga ke tingkat masyarakat paling bawah. Kegiatan sosialisasi ini menjadi instrumen pengawasan preventif, di mana potensi permasalahan implementasi kebijakan dapat diidentifikasi sejak dini melalui dialog langsung dengan pemangku kepentingan di daerah.

Malang Raya, yang mencakup Kota Malang, Kota Batu, dan Kabupaten Malang, dipilih sebagai lokasi kegiatan mengingat wilayah ini merupakan basis konstituen Dapil Jawa Timur V yang memiliki karakteristik ekonomi yang beragam, mulai dari sektor pertanian, perkebunan, perdagangan, hingga usaha mikro, kecil, dan menengah. Keberagaman ini menuntut pendekatan sosialisasi yang komprehensif dan melibatkan berbagai unsur, termasuk pemerintah daerah, dinas teknis terkait, asosiasi usaha, kelompok masyarakat, hingga aparat pengawas internal dan eksternal, agar substansi regulasi dapat dipahami secara utuh oleh seluruh lapisan masyarakat yang terdampak, baik secara langsung maupun tidak langsung.

Dalam praktiknya, tantangan utama dalam sosialisasi produk perundang-undangan di tingkat daerah adalah kesenjangan informasi antara regulasi yang diterbitkan di tingkat pusat dengan pemahaman aktual di lapangan. Tidak jarang ditemukan bahwa suatu kebijakan yang secara normatif sudah dirancang dengan baik, dalam pelaksanaannya mengalami hambatan akibat minimnya sosialisasi, kurangnya koordinasi antar instansi, maupun keterbatasan kapasitas sumber daya manusia di daerah. Kunjungan kerja ini diharapkan menjembatani kesenjangan tersebut, sekaligus menjadi sarana menyerap masukan dan evaluasi langsung dari masyarakat.

Dengan semangat tersebut, kegiatan sosialisasi undang-undang ini dilaksanakan dengan tujuan utama memberikan pemahaman yang komprehensif kepada masyarakat dan pemangku kepentingan di Malang Raya mengenai substansi regulasi yang disosialisasikan, sekaligus membuka ruang dialog untuk menampung aspirasi dan masukan konstruktif bagi evaluasi kebijakan di tingkat nasional."""

for paragraph in pendahuluan_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
add_centered("Malang, 11 April 2025")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("Nomor Anggota A-220")

doc.add_page_break()

# ==================== ISI KEGIATAN ====================
add_title("2. ISI KEGIATAN", 14)

photo_folder = r"D:\DATA PAPA\LAPORAN\LAPSOS UU\LAPSOS 1"

isi_text = """Dalam upaya memperkuat ketahanan pangan nasional dan meningkatkan kesejahteraan petani, Peraturan Presiden (Perpres) Nomor 6 Tahun 2025 tentang Tata Kelola Pupuk Bersubsidi telah ditetapkan sebagai regulasi penting yang menjawab tantangan distribusi pupuk bersubsidi di berbagai daerah. Sebagai bagian dari komitmen negara dalam memastikan kebijakan ini berjalan secara transparan, tepat sasaran, dan akuntabel, Ketua Badan Akuntabilitas Keuangan Negara (BAKN) DPR RI sekaligus anggota Komisi XI DPR RI dari Fraksi PDI Perjuangan, Bapak Ir. Andreas Eddy Susetyo, M.M., hadir langsung dalam kegiatan sosialisasi dan pemantauan alur distribusi pupuk bersubsidi.

Kegiatan ini dilaksanakan di Kantor Dinas Tanaman Pangan, Hortikultura, dan Perkebunan yang berlokasi di Jalan Sumedang No. 28, Cokolelo, Kepanjen, Kecamatan Kepanjen, Kabupaten Malang, Jawa Timur. Dihadiri oleh para stakeholder yang terkait dalam tata kelola pupuk bersubsidi, seperti perwakilan dari dinas pertanian, distributor, pengecer, kelompok tani, serta aparat pengawas internal dan eksternal.

Dalam sambutannya, Bapak Andreas menegaskan bahwa Perpres ini lahir dari kebutuhan riil di lapangan yang menuntut adanya sistem distribusi pupuk yang lebih tertib, adil, dan mengedepankan prinsip transparansi. "Kita harus memastikan bahwa subsidi pupuk benar-benar sampai ke tangan petani yang berhak, bukan disalahgunakan oleh oknum tidak bertanggung jawab," ujar beliau."""

for paragraph in isi_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "L1.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "L2.JPG"), width=Inches(5.5))

doc.add_paragraph()

lanjutan_text = """Beliau juga menekankan pentingnya sinergi antar instansi dan penguatan pengawasan berlapis untuk mencegah kebocoran anggaran subsidi. Dengan pengawasan digital melalui sistem e-RDKK dan integrasi data petani secara nasional, pemerintah berkomitmen memperbaiki tata kelola distribusi pupuk agar lebih efisien dan berkeadilan.

Sebagai wakil rakyat dari Malang Raya, beliau menutup sambutannya dengan ajakan kepada seluruh pemangku kepentingan untuk bersama-sama mengawal implementasi Perpres ini demi keberlangsungan pertanian dan kesejahteraan petani di daerah. "Kami di DPR RI, khususnya di Komisi XI dan BAKN, akan terus mengawal agar setiap rupiah subsidi negara memberi manfaat optimal bagi rakyat, khususnya para petani," pungkasnya."""

for paragraph in lanjutan_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "L3.JPG"), width=Inches(5.5))

doc.add_page_break()

# ==================== PENUTUP ====================
add_title("3. PENUTUP", 14)

penutup_text = """Kegiatan sosialisasi dan pemantauan alur distribusi pupuk bersubsidi ini menjadi momentum penting dalam memastikan implementasi Perpres Nomor 6 Tahun 2025 berjalan dengan baik dan berpihak pada kepentingan petani. Diharapkan seluruh stakeholder yang hadir dapat memahami secara mendalam substansi regulasi ini dan turut aktif mengawasi pelaksanaannya di lapangan.

Peran serta masyarakat, kelompok tani, dan aparatur pemerintah daerah menjadi kunci suksesnya tata kelola pupuk bersubsidi yang transparan dan tepat sasaran. Dengan kolaborasi yang kuat, segala bentuk penyimpangan dapat diminimalisir, sehingga pupuk bersubsidi benar-benar menjadi alat bantu produktivitas petani.

Melalui regulasi yang kuat dan pengawasan yang berkesinambungan, kita berharap sektor pertanian di Indonesia semakin maju, berdaya saing, dan mampu meningkatkan kesejahteraan petani secara menyeluruh. Mari kawal bersama demi masa depan pertanian yang berkelanjutan.

Dari rangkaian kegiatan sosialisasi yang telah dilaksanakan, dapat disimpulkan bahwa keberhasilan sebuah regulasi tidak hanya ditentukan oleh kualitas substansi hukumnya, tetapi juga sangat bergantung pada efektivitas proses sosialisasi dan pemahaman para pemangku kepentingan di lapangan. Dialog langsung antara Anggota DPR RI dengan dinas teknis, distributor, pengecer, kelompok tani, dan aparat pengawas telah membuka ruang komunikasi yang konstruktif, di mana berbagai kendala teknis maupun administratif yang selama ini menjadi hambatan implementasi dapat diidentifikasi secara lebih akurat dan komprehensif.

Sebagai Ketua Badan Akuntabilitas Keuangan Negara sekaligus Anggota Komisi XI DPR RI, komitmen untuk terus mengawal implementasi kebijakan yang berkaitan dengan pengelolaan keuangan dan subsidi negara akan terus dijalankan secara berkesinambungan. Hasil dari kegiatan sosialisasi ini akan menjadi bahan masukan yang berharga bagi pembahasan lebih lanjut di tingkat Komisi XI maupun BAKN, khususnya dalam rangka evaluasi kebijakan subsidi dan penguatan pengawasan anggaran negara agar setiap rupiah yang dianggarkan benar-benar memberikan manfaat optimal bagi masyarakat yang berhak menerimanya.

Selain itu, ditemukan pula bahwa masih terdapat kesenjangan pemahaman di sebagian kecil pemangku kepentingan di tingkat lapangan mengenai mekanisme teknis pelaksanaan regulasi, khususnya terkait sistem digital seperti e-RDKK dan integrasi data penerima manfaat. Hal ini menjadi catatan penting yang perlu ditindaklanjuti melalui sosialisasi lanjutan maupun pendampingan teknis yang lebih intensif dari instansi terkait, agar seluruh pemangku kepentingan dapat mengoperasikan sistem tersebut secara optimal dan tidak terjadi hambatan administratif yang merugikan penerima manfaat di lapangan.

Ke depan, kegiatan sosialisasi undang-undang semacam ini perlu terus dilakukan secara berkala dan menjangkau lebih banyak wilayah di Dapil Jawa Timur V, mengingat pentingnya memastikan bahwa setiap produk perundang-undangan yang dihasilkan oleh DPR RI benar-benar dipahami dan dapat diimplementasikan dengan baik oleh masyarakat di tingkat akar rumput. Dukungan dan partisipasi aktif dari seluruh elemen masyarakat, pemerintah daerah, serta sektor swasta akan menjadi kunci utama keberhasilan implementasi kebijakan secara menyeluruh dan berkelanjutan, demi terwujudnya tata kelola pemerintahan yang transparan, akuntabel, dan berpihak pada kepentingan rakyat.

Akhir kata, kegiatan Sosialisasi Undang-Undang Ke-1 ini diharapkan dapat menjadi model pelaksanaan tugas konstitusional yang efektif bagi Anggota DPR RI dalam menjembatani kebijakan pusat dengan kebutuhan riil masyarakat di daerah. Semoga hasil dari kegiatan ini dapat memberikan kontribusi nyata bagi perbaikan tata kelola kebijakan publik dan peningkatan kesejahteraan masyarakat Malang Raya secara berkelanjutan."""

for paragraph in penutup_text.split('\n\n'):
    if paragraph.strip():
        add_normal(paragraph.strip())

doc.add_paragraph()
add_photo(os.path.join(photo_folder, "L4.JPG"), width=Inches(5.5))
doc.add_paragraph()
add_photo(os.path.join(photo_folder, "L5.JPG"), width=Inches(5.5))

doc.add_paragraph()
doc.add_paragraph()

add_centered("Malang, 11 April 2025")
doc.add_paragraph()
add_centered("Ir. Andreas Eddy Susetyo, MM")
add_centered("Nomor Anggota A-220")

# Save document
output_path = r"D:\DATA PAPA\LAPORAN\LAPSOS UU\LAPSOS 1\laporan sos UU ke 1 tgl 9-10 April 2025.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
