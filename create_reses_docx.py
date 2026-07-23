"""
Script untuk membuat Laporan Kunjungan Kerja Reses Perseorangan DPR RI
Masa Sidang V Tahun Sidang 2025-2026 - Ir. Andreas Eddy Susetyo, MM (A-220)
Periode Reses: 22 Juli - 13 Agustus 2026
Daerah Pemilihan Jawa Timur V (Kota Malang, Kabupaten Malang, Kota Batu)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# CONFIG
# ============================================================
OUTPUT_DIR = r"D:\DATA PAPA\LAPORAN\RESES\JULI_AGUSTUS 2026"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "FINAL_LAPORAN_RESES_Andreas_Eddy_Susetyo_Juli_Agustus_2026.docx")
FOTO_DIR = r"D:\DATA PAPA\LAPORAN\RESES\JULI_AGUSTUS 2026"

MASA_SIDANG = "MASA SIDANG V TAHUN SIDANG 2025-2026"
PERIODE_RESES = "22 Juli - 13 Agustus 2026"
DAPIL = "Jawa Timur V (Kota Malang, Kabupaten Malang, Kota Batu)"
TANGGAL_LAPORAN = "Malang, 15 Agustus 2026"

# ============================================================
# DATA KEGIATAN (27 Juli - 4 Agustus 2026)
# ============================================================
KEGIATAN = [
    {
        "no": 1, "hari_tanggal": "Senin, 27 Juli 2026", "waktu": "09.00-11.30 WIB",
        "tempat": "Kecamatan Jabung, Kabupaten Malang",
        "judul": "Pertemuan dan Dialog dengan Pengurus serta Anggota Koperasi",
        "usulan": [
            "Diusulkan kemudahan akses permodalan bagi koperasi melalui skema bunga rendah dan penyederhanaan persyaratan pengajuan pinjaman.",
            "Berharap adanya pelatihan manajemen keuangan dan digitalisasi tata kelola koperasi agar lebih transparan dan akuntabel.",
            "Mendesak pemerintah memberikan insentif fiskal bagi koperasi yang aktif menyalurkan modal usaha kepada anggota di sektor pertanian dan perdagangan.",
        ],
        "fotos": ["R1.JPG", "R2.JPG"],
        "deskripsi_foto": (
            "Pertemuan dan dialog dengan pengurus serta anggota Koperasi di Kecamatan Jabung, "
            "Kabupaten Malang, berlangsung pada Senin pagi dalam suasana yang akrab dan terbuka. "
            "Para pengurus koperasi menyampaikan berbagai kendala yang dihadapi dalam pengelolaan "
            "modal usaha, terutama keterbatasan akses pembiayaan dengan bunga terjangkau bagi anggota "
            "yang mayoritas bergerak di sektor pertanian dan perdagangan kecil. Dialog turut membahas "
            "pentingnya digitalisasi administrasi koperasi agar pengelolaan keuangan lebih transparan "
            "dan mudah diawasi. Pertemuan ini menjadi masukan berharga mengenai kebutuhan riil "
            "penguatan kelembagaan ekonomi rakyat di tingkat kecamatan."
        ),
    },
    {
        "no": 2, "hari_tanggal": "Senin, 27 Juli 2026", "waktu": "12.30-14.30 WIB",
        "tempat": "Kota Malang",
        "judul": "Dialog dengan Pelaku Industri Jasa Keuangan Malang Raya",
        "usulan": [
            "Mengharapkan penguatan literasi keuangan masyarakat guna menekan angka pinjaman daring ilegal yang marak di Malang Raya.",
            "Diusulkan insentif dan regulasi yang mendorong lembaga keuangan formal memperluas jangkauan kredit bagi pelaku UMKM.",
            "Mendesak pengawasan yang lebih ketat dari OJK terhadap praktik fintech lending guna melindungi konsumen dari suku bunga tidak wajar.",
        ],
        "fotos": ["R3.JPG", "R4.JPG"],
        "deskripsi_foto": (
            "Dialog dengan para pelaku industri jasa keuangan Malang Raya di Kota Malang membahas "
            "dinamika sektor keuangan formal dan informal yang berkembang di tengah masyarakat. "
            "Para pelaku industri menyoroti maraknya praktik pinjaman daring ilegal yang menjerat "
            "masyarakat dengan bunga tidak wajar akibat rendahnya literasi keuangan. Dibahas pula "
            "peluang perluasan kredit UMKM melalui lembaga keuangan formal yang lebih inklusif serta "
            "perlunya pengawasan yang lebih ketat terhadap penyelenggara fintech lending. Pertemuan "
            "ini menegaskan pentingnya edukasi keuangan sebagai benteng perlindungan masyarakat."
        ),
    },
    {
        "no": 3, "hari_tanggal": "Senin, 27 Juli 2026", "waktu": "16.00-17.30 WIB",
        "tempat": "Kecamatan Kedungkandang, Kota Malang",
        "judul": "Dialog Serap Aspirasi dengan Orang Tua dan Siswa SMA/SMK",
        "usulan": [
            "Meminta peningkatan sarana dan prasarana pendidikan, terutama laboratorium dan peralatan praktik SMK guna mendukung kesiapan kerja lulusan.",
            "Berharap program beasiswa dan bantuan biaya pendidikan bagi siswa dari keluarga kurang mampu agar tidak putus sekolah.",
            "Diusulkan penguatan kerja sama sekolah dengan dunia usaha dan dunia industri (DUDI) untuk memperluas peluang magang dan penyerapan tenaga kerja lulusan SMK.",
        ],
        "fotos": ["R5.JPG", "R6.JPG"],
        "deskripsi_foto": (
            "Pertemuan dan dialog serap aspirasi dengan orang tua dan siswa SMA/SMK di Kecamatan "
            "Kedungkandang, Kota Malang, berlangsung dalam suasana penuh semangat menjelang tahun "
            "ajaran baru. Para orang tua menyampaikan keluhan mengenai biaya pendidikan yang masih "
            "membebani keluarga berpenghasilan rendah, sementara pihak siswa SMK menyoroti minimnya "
            "peralatan praktik yang memadai untuk mendukung kompetensi keahlian. Dialog juga membahas "
            "pentingnya jalinan kerja sama sekolah dengan dunia usaha agar lulusan SMK lebih siap "
            "terserap ke dunia kerja. Aspirasi ini menjadi catatan penting bagi peningkatan mutu "
            "pendidikan vokasi di wilayah Kota Malang."
        ),
    },
    {
        "no": 4, "hari_tanggal": "Selasa, 28 Juli 2026", "waktu": "08.30-11.30 WIB",
        "tempat": "Paroki Kota Batu",
        "judul": "Menghadiri Kegiatan Rutin Umat Katolik Paroki Kota Batu",
        "usulan": [
            "Berharap dukungan pemerintah dalam pemeliharaan dan pengembangan fasilitas sarana ibadah serta kegiatan sosial keagamaan.",
            "Diusulkan penguatan program kerukunan antarumat beragama melalui kegiatan lintas iman yang melibatkan generasi muda.",
            "Meminta perhatian terhadap program sosial gereja seperti pelayanan kesehatan dan pendidikan bagi jemaat kurang mampu.",
        ],
        "fotos": ["R7.JPG", "R8.JPG"],
        "deskripsi_foto": (
            "Menghadiri kegiatan rutin bersama Umat Katolik di Paroki Kota Batu menjadi bagian dari "
            "upaya menjaga kedekatan dengan seluruh elemen masyarakat lintas keyakinan di Daerah "
            "Pemilihan Jawa Timur V. Pengurus paroki dan jemaat menyampaikan apresiasi atas kehadiran "
            "wakil rakyat serta menyampaikan harapan agar pemerintah turut memperhatikan pemeliharaan "
            "fasilitas ibadah dan program sosial gereja yang selama ini menjadi penopang bagi jemaat "
            "kurang mampu. Pertemuan ini juga memperkuat semangat kerukunan dan toleransi antarumat "
            "beragama di wilayah Kota Batu."
        ),
    },
    {
        "no": 5, "hari_tanggal": "Selasa, 28 Juli 2026", "waktu": "13.00-15.30 WIB",
        "tempat": "Kecamatan Klojen, Kota Malang",
        "judul": "Kunjungan ke Kantor Perum Perhutani KPH Malang",
        "usulan": [
            "Mengusulkan penguatan program perhutanan sosial agar masyarakat sekitar hutan memperoleh akses legal pemanfaatan lahan untuk penghidupan.",
            "Mendesak peningkatan koordinasi antara Perhutani, pemerintah daerah, dan masyarakat dalam pencegahan kebakaran hutan dan lahan.",
            "Berharap adanya program kemitraan ekonomi kehutanan yang melibatkan kelompok tani hutan guna meningkatkan kesejahteraan masyarakat sekitar kawasan hutan.",
        ],
        "fotos": ["R9.jpg", "R10.jpg"],
        "deskripsi_foto": (
            "Kunjungan ke Kantor Perum Perhutani Kesatuan Pemangkuan Hutan (KPH) Malang di Kecamatan "
            "Klojen, Kota Malang, membahas berbagai isu pengelolaan kawasan hutan yang berdampak "
            "langsung pada masyarakat sekitar. Pihak Perhutani memaparkan program perhutanan sosial "
            "yang telah berjalan namun masih perlu diperluas cakupannya agar lebih banyak masyarakat "
            "memperoleh akses legal pemanfaatan lahan hutan. Dibahas pula pentingnya sinergi lintas "
            "pihak dalam mencegah kebakaran hutan dan lahan pada musim kemarau, serta peluang "
            "kemitraan ekonomi yang dapat meningkatkan kesejahteraan kelompok tani hutan."
        ),
    },
    {
        "no": 6, "hari_tanggal": "Rabu, 29 Juli 2026", "waktu": "09.00-12.00 WIB",
        "tempat": "Sekretariat DPC Kota Malang, Jl. Raden Panji Suroso, Kec. Blimbing, Kota Malang",
        "judul": "Menghadiri Konsolidasi DPC Partai PDI Perjuangan Kota Malang",
        "usulan": [
            "Diusulkan penguatan konsolidasi struktur partai hingga tingkat ranting untuk mengoptimalkan penyerapan aspirasi masyarakat.",
            "Mengharapkan sinergi program kerja partai dengan program pemerintah daerah agar aspirasi konstituen dapat ditindaklanjuti secara konkret.",
            "Meminta peningkatan kaderisasi dan pelatihan politik bagi kader muda partai di Kota Malang.",
        ],
        "fotos": ["R11.JPG", "R12.JPG"],
        "deskripsi_foto": (
            "Menghadiri kegiatan konsolidasi Dewan Pimpinan Cabang (DPC) Partai PDI Perjuangan Kota "
            "Malang di Sekretariat DPC, Jalan Raden Panji Suroso, Kecamatan Blimbing, menjadi momentum "
            "penting untuk memperkuat soliditas struktur partai di tingkat kota. Para pengurus dan "
            "kader menyampaikan evaluasi kinerja penyerapan aspirasi di tingkat ranting serta harapan "
            "agar program partai dapat lebih sinergis dengan kebijakan pemerintah daerah. Pertemuan "
            "ini juga membahas pentingnya kaderisasi politik generasi muda sebagai bagian dari "
            "regenerasi kepemimpinan partai di Kota Malang."
        ),
    },
    {
        "no": 7, "hari_tanggal": "Rabu, 29 Juli 2026", "waktu": "13.00-15.30 WIB",
        "tempat": "Kepanjen, Kabupaten Malang",
        "judul": "Dialog dan Serap Aspirasi dengan Pelaku BPR Kabupaten Malang",
        "usulan": [
            "Mendesak relaksasi kebijakan permodalan bagi BPR agar mampu bersaing dengan lembaga keuangan digital dan perbankan nasional.",
            "Diusulkan penyederhanaan regulasi OJK terkait BPR agar biaya kepatuhan tidak membebani operasional lembaga keuangan skala kecil.",
            "Berharap dukungan pemerintah daerah dalam memperluas kerja sama BPR dengan program pembiayaan UMKM dan sektor pertanian.",
        ],
        "fotos": ["R13.JPG", "R14.JPG"],
        "deskripsi_foto": (
            "Dialog dan serap aspirasi dengan para pelaku Bank Perkreditan Rakyat (BPR) Kabupaten "
            "Malang di Kepanjen membahas tantangan yang dihadapi lembaga keuangan skala kecil di "
            "tengah persaingan dengan perbankan digital dan lembaga keuangan nasional. Para pelaku "
            "BPR menyampaikan keluhan mengenai beratnya biaya kepatuhan regulasi yang membebani "
            "operasional, serta perlunya relaksasi kebijakan permodalan agar BPR dapat terus "
            "berkontribusi dalam pembiayaan UMKM dan sektor pertanian di Kabupaten Malang. Dialog ini "
            "menjadi masukan penting bagi penguatan lembaga keuangan mikro daerah."
        ),
    },
    {
        "no": 8, "hari_tanggal": "Kamis, 30 Juli 2026", "waktu": "09.00-11.30 WIB",
        "tempat": "Kecamatan Wonosari, Kabupaten Malang",
        "judul": "Kunjungan ke Koperasi Simpan Pinjam",
        "usulan": [
            "Meminta pendampingan teknis dan pelatihan manajemen risiko kredit bagi pengurus koperasi simpan pinjam.",
            "Mendesak perlindungan hukum bagi anggota koperasi dari praktik simpan pinjam ilegal yang berkedok koperasi.",
            "Diusulkan tambahan akses permodalan dari lembaga keuangan pemerintah untuk memperkuat kapasitas penyaluran pinjaman koperasi.",
        ],
        "fotos": ["R15.JPG", "R16.JPG"],
        "deskripsi_foto": (
            "Kunjungan ke Koperasi Simpan Pinjam di Kecamatan Wonosari, Kabupaten Malang, menjadi "
            "ajang dialog langsung dengan pengurus mengenai pengelolaan simpan pinjam bagi anggota "
            "koperasi yang mayoritas warga sekitar. Pengurus koperasi menyampaikan kebutuhan "
            "pendampingan teknis dalam manajemen risiko kredit agar praktik simpan pinjam tetap sehat "
            "dan berkelanjutan. Turut disampaikan keresahan atas maraknya lembaga simpan pinjam "
            "ilegal yang berkedok koperasi dan merugikan masyarakat, sehingga diperlukan perlindungan "
            "hukum serta dukungan permodalan tambahan dari lembaga keuangan pemerintah."
        ),
    },
    {
        "no": 9, "hari_tanggal": "Kamis, 30 Juli 2026", "waktu": "12.30-14.30 WIB",
        "tempat": "Kota Malang",
        "judul": "Menghadiri Misa Perayaan Ulang Tahun Gereja",
        "usulan": [
            "Berharap dukungan pemerintah terhadap pemeliharaan bangunan gereja bersejarah sebagai bagian dari cagar budaya.",
            "Mengusulkan penguatan program toleransi dan kerukunan umat beragama di lingkungan perkotaan.",
            "Meminta perhatian terhadap keamanan dan ketertiban penyelenggaraan kegiatan keagamaan berskala besar.",
        ],
        "fotos": ["R17.JPG", "R18.JPG"],
        "deskripsi_foto": (
            "Menghadiri misa perayaan Ulang Tahun Gereja di Kota Malang menjadi kesempatan untuk "
            "menyampaikan apresiasi atas peran gereja dalam menjaga kehidupan sosial dan spiritual "
            "umat. Pengurus gereja menyampaikan harapan agar pemerintah turut memperhatikan "
            "pemeliharaan bangunan gereja yang memiliki nilai sejarah sebagai bagian dari cagar "
            "budaya kota. Perayaan ini juga menjadi momen mempererat kerukunan antarumat beragama di "
            "lingkungan perkotaan, sekaligus menyoroti pentingnya dukungan keamanan dalam "
            "penyelenggaraan kegiatan keagamaan berskala besar."
        ),
    },
    {
        "no": 10, "hari_tanggal": "Kamis, 30 Juli 2026", "waktu": "15.30-16.30 WIB",
        "tempat": "Kelurahan Turen, Kecamatan Turen, Kabupaten Malang",
        "judul": "Kunjungan dan Dialog Serap Aspirasi dengan Warga Kelurahan Turen",
        "usulan": [
            "Mendesak perbaikan infrastruktur jalan dan drainase lingkungan yang rusak akibat curah hujan tinggi.",
            "Berharap peningkatan layanan kesehatan dasar melalui optimalisasi fasilitas Puskesmas setempat.",
            "Meminta program bantuan sosial yang tepat sasaran bagi warga terdampak pelemahan daya beli.",
        ],
        "fotos": ["R19.JPG", "R20.JPG"],
        "deskripsi_foto": (
            "Kunjungan dan dialog serap aspirasi dengan warga di Kelurahan Turen, Kecamatan Turen, "
            "Kabupaten Malang, mengungkap berbagai persoalan yang dihadapi masyarakat sehari-hari. "
            "Warga menyampaikan keluhan mengenai kondisi jalan dan saluran drainase yang rusak dan "
            "kerap menyebabkan genangan saat curah hujan tinggi. Selain itu, warga juga berharap "
            "adanya peningkatan layanan kesehatan dasar di Puskesmas setempat serta penyaluran "
            "bantuan sosial yang lebih tepat sasaran di tengah tekanan ekonomi yang dirasakan "
            "masyarakat."
        ),
    },
    {
        "no": 11, "hari_tanggal": "Jumat, 31 Juli 2026", "waktu": "08.30-11.00 WIB",
        "tempat": "Kecamatan Sumberpucung, Kabupaten Malang",
        "judul": "Serap Aspirasi dengan Perempuan Pelaku UMKM",
        "usulan": [
            "Diusulkan pelatihan digital marketing dan akses pembiayaan mikro khusus bagi perempuan pelaku usaha.",
            "Meminta program pendampingan perizinan usaha seperti NIB dan sertifikasi halal yang mudah diakses oleh UMKM perempuan.",
            "Berharap fasilitasi pemasaran produk UMKM perempuan melalui pameran dan kemitraan dengan ritel modern.",
        ],
        "fotos": ["R21.JPG", "R22.JPG"],
        "deskripsi_foto": (
            "Serap aspirasi dengan perempuan pelaku UMKM di Kecamatan Sumberpucung, Kabupaten Malang, "
            "menghadirkan diskusi hangat mengenai tantangan usaha yang digeluti para ibu rumah tangga "
            "dan perempuan wirausaha. Para pelaku UMKM menyampaikan kebutuhan pelatihan pemasaran "
            "digital serta akses pembiayaan mikro yang selama ini sulit dijangkau. Diusulkan pula "
            "adanya pendampingan pengurusan izin usaha seperti Nomor Induk Berusaha dan sertifikasi "
            "halal, serta fasilitasi pemasaran produk melalui pameran dan kemitraan dengan jaringan "
            "ritel modern agar produk UMKM perempuan semakin dikenal luas."
        ),
    },
    {
        "no": 12, "hari_tanggal": "Jumat, 31 Juli 2026", "waktu": "15.30-16.30 WIB",
        "tempat": "Desa Druju, Kecamatan Sumbermanjing Wetan, Kabupaten Malang",
        "judul": "Serap Aspirasi dan Kunjungan Lapangan ke Desa Druju",
        "usulan": [
            "Mendesak perbaikan akses jalan menuju desa yang selama ini menghambat distribusi hasil pertanian dan perkebunan.",
            "Meminta perluasan jaringan listrik dan telekomunikasi di wilayah pedesaan yang masih terbatas.",
            "Diusulkan program peningkatan produktivitas pertanian melalui bantuan bibit unggul dan pupuk bersubsidi.",
        ],
        "fotos": ["R23.JPG", "R24.JPG"],
        "deskripsi_foto": (
            "Serap aspirasi dan kunjungan lapangan ke Desa Druju, Kecamatan Sumbermanjing Wetan, "
            "Kabupaten Malang, memberikan gambaran nyata mengenai keterbatasan infrastruktur di "
            "wilayah selatan Malang. Warga desa menyampaikan bahwa akses jalan yang rusak menghambat "
            "distribusi hasil pertanian dan perkebunan ke pasar. Selain itu, jaringan listrik dan "
            "telekomunikasi di sejumlah dusun masih terbatas sehingga menghambat aktivitas ekonomi dan "
            "pendidikan warga. Kunjungan ini menegaskan urgensi pemerataan pembangunan infrastruktur "
            "dasar di wilayah pedesaan."
        ),
    },
    {
        "no": 13, "hari_tanggal": "Sabtu, 1 Agustus 2026", "waktu": "08.30-11.30 WIB",
        "tempat": "Desa Tumpang, Kecamatan Tumpang, Kabupaten Malang",
        "judul": "Pertemuan dengan Pengurus PAC Partai PDI Perjuangan Kabupaten Malang",
        "usulan": [
            "Diusulkan penguatan koordinasi antara PAC dan Ranting dalam menjaring aspirasi masyarakat di tingkat kecamatan.",
            "Berharap dukungan program pemberdayaan ekonomi masyarakat berbasis potensi lokal Kecamatan Tumpang, khususnya sektor pariwisata.",
            "Meminta peningkatan kaderisasi politik generasi muda sebagai bagian dari regenerasi kepemimpinan partai di daerah.",
        ],
        "fotos": ["R25.jpg", "R26.jpg"],
        "deskripsi_foto": (
            "Pertemuan dengan sejumlah pengurus Pimpinan Anak Cabang (PAC) Partai PDI Perjuangan "
            "Kabupaten Malang di Desa Tumpang, Kecamatan Tumpang, membahas penguatan konsolidasi "
            "partai hingga tingkat ranting. Para pengurus PAC menyampaikan pentingnya koordinasi yang "
            "lebih intens dalam menjaring aspirasi masyarakat, termasuk dukungan terhadap "
            "pengembangan potensi wisata lokal Kecamatan Tumpang sebagai sumber penggerak ekonomi "
            "warga. Pertemuan ini juga menyoroti perlunya kaderisasi politik generasi muda sebagai "
            "bagian dari regenerasi kepemimpinan partai di tingkat kecamatan."
        ),
    },
    {
        "no": 14, "hari_tanggal": "Sabtu, 1 Agustus 2026", "waktu": "13.00-15.30 WIB",
        "tempat": "Desa Sitiarjo, Kecamatan Sumbermanjing Wetan, Kabupaten Malang",
        "judul": "Serap Aspirasi dengan Petani dan Nelayan Sumbermanjing Wetan",
        "usulan": [
            "Mendesak jaminan harga hasil pertanian dan tangkapan nelayan yang stabil melalui intervensi pasar oleh pemerintah.",
            "Meminta bantuan sarana produksi seperti alat tangkap ikan dan pupuk bersubsidi bagi petani dan nelayan pesisir selatan.",
            "Berharap pembangunan infrastruktur pendukung seperti tempat pelelangan ikan (TPI) dan akses jalan menuju kawasan pesisir.",
        ],
        "fotos": ["R27.JPG", "R28.JPG"],
        "deskripsi_foto": (
            "Serap aspirasi dengan petani dan nelayan di Kecamatan Sumbermanjing Wetan, bertempat di "
            "Desa Sitiarjo, mengangkat persoalan ganda yang dihadapi masyarakat pesisir selatan "
            "Malang. Para petani mengeluhkan fluktuasi harga hasil panen yang tidak menentu, "
            "sementara nelayan menyampaikan keterbatasan sarana tangkap dan minimnya infrastruktur "
            "pendukung seperti tempat pelelangan ikan. Kedua kelompok berharap pemerintah hadir "
            "melalui intervensi pasar yang menjaga stabilitas harga serta bantuan sarana produksi "
            "yang memadai bagi keberlangsungan usaha tani dan tangkap di wilayah pesisir."
        ),
    },
    {
        "no": 15, "hari_tanggal": "Minggu, 2 Agustus 2026", "waktu": "09.00-11.30 WIB",
        "tempat": "Desa Sambigede, Kecamatan Sumberpucung, Kabupaten Malang",
        "judul": "Kunjungan dan Dialog dengan Warga Desa Sambigede",
        "usulan": [
            "Diusulkan perbaikan sistem irigasi pertanian guna mendukung produktivitas lahan sawah warga.",
            "Berharap peningkatan akses air bersih bagi warga yang masih mengandalkan sumber air terbatas pada musim kemarau.",
            "Meminta program padat karya untuk menyerap tenaga kerja lokal di sektor infrastruktur desa.",
        ],
        "fotos": ["R29.JPG", "R30.JPG"],
        "deskripsi_foto": (
            "Kunjungan dan dialog dengan warga di Desa Sambigede, Kecamatan Sumberpucung, Kabupaten "
            "Malang, menyoroti kebutuhan mendasar masyarakat agraris di wilayah tersebut. Warga "
            "menyampaikan bahwa sistem irigasi yang ada belum optimal sehingga berdampak pada "
            "produktivitas lahan sawah, terutama pada musim kemarau. Ketersediaan air bersih untuk "
            "kebutuhan sehari-hari juga menjadi keluhan utama warga. Dialog ini turut membahas usulan "
            "program padat karya sebagai upaya menyerap tenaga kerja lokal sekaligus memperbaiki "
            "infrastruktur dasar desa."
        ),
    },
    {
        "no": 16, "hari_tanggal": "Minggu, 2 Agustus 2026", "waktu": "12.30-14.30 WIB",
        "tempat": "Desa Kalipare, Kecamatan Kalipare, Kabupaten Malang",
        "judul": "Kunjungan dan Dialog dengan Warga Pelaku UMKM Desa Kalipare",
        "usulan": [
            "Meminta akses permodalan KUR yang lebih mudah dijangkau oleh pelaku UMKM di wilayah pedesaan.",
            "Diusulkan pelatihan pengemasan dan standardisasi produk agar mampu bersaing di pasar yang lebih luas.",
            "Mendesak perbaikan akses jalan produksi guna memperlancar distribusi hasil UMKM ke luar wilayah kecamatan.",
        ],
        "fotos": ["R31.JPG", "R32.JPG"],
        "deskripsi_foto": (
            "Kunjungan dan dialog dengan warga pelaku UMKM di Desa Kalipare, Kecamatan Kalipare, "
            "Kabupaten Malang, mengungkap potensi ekonomi lokal yang masih menghadapi kendala akses "
            "permodalan. Para pelaku UMKM menyampaikan kesulitan menjangkau skema Kredit Usaha Rakyat "
            "(KUR) yang persyaratannya dinilai masih rumit bagi usaha skala kecil di pedesaan. "
            "Diusulkan pula pelatihan pengemasan dan standardisasi produk agar mampu bersaing di "
            "pasar yang lebih luas, serta perbaikan akses jalan produksi guna memperlancar distribusi "
            "hasil UMKM keluar wilayah kecamatan."
        ),
    },
    {
        "no": 17, "hari_tanggal": "Senin, 3 Agustus 2026", "waktu": "08.30-11.30 WIB",
        "tempat": "Sekretariat DPC Kota Batu, Jl. Oro-Oro Ombo, Kota Batu",
        "judul": "Kunjungan dan Konsolidasi dengan Jajaran Pengurus DPC Kota Batu",
        "usulan": [
            "Diusulkan penguatan konsolidasi kader di tingkat kelurahan guna memperkuat basis dukungan partai di Kota Batu.",
            "Berharap sinkronisasi program partai dengan agenda pembangunan pariwisata Kota Batu yang menjadi sektor unggulan.",
            "Meminta peningkatan kapasitas kader dalam pengawasan pelaksanaan program pemerintah daerah.",
        ],
        "fotos": ["R33.JPG", "R34.JPG"],
        "deskripsi_foto": (
            "Kunjungan dan konsolidasi dengan jajaran pengurus DPC Kota Batu di Sekretariat DPC, Jalan "
            "Oro-Oro Ombo, membahas penguatan struktur organisasi partai di tingkat kelurahan. Para "
            "pengurus menyampaikan evaluasi kinerja kader dalam menjaring aspirasi masyarakat serta "
            "usulan agar program partai lebih sinkron dengan agenda pembangunan pariwisata yang "
            "menjadi sektor unggulan Kota Batu. Konsolidasi ini juga menekankan pentingnya peningkatan "
            "kapasitas kader dalam mengawasi pelaksanaan program pemerintah daerah agar tepat sasaran."
        ),
    },
    {
        "no": 18, "hari_tanggal": "Senin, 3 Agustus 2026", "waktu": "15.30-16.30 WIB",
        "tempat": "Desa Sumbersekar, Kecamatan Dau, Kabupaten Malang",
        "judul": "Menghadiri Kerja Bakti di Desa Sumbersekar",
        "usulan": [
            "Meminta dukungan anggaran untuk perbaikan fasilitas umum desa seperti saluran air dan jalan lingkungan.",
            "Berharap penguatan program gotong royong sebagai bagian dari pemberdayaan masyarakat berbasis kearifan lokal.",
            "Mendesak perhatian terhadap pengelolaan sampah dan kebersihan lingkungan di kawasan permukiman padat penduduk.",
        ],
        "fotos": ["R35.JPG", "R36.JPG"],
        "deskripsi_foto": (
            "Menghadiri kegiatan kerja bakti bersama warga Desa Sumbersekar, Kecamatan Dau, Kabupaten "
            "Malang, menjadi wujud kedekatan langsung dengan aktivitas gotong royong masyarakat. Warga "
            "menyampaikan kebutuhan dukungan anggaran untuk perbaikan fasilitas umum seperti saluran "
            "air dan jalan lingkungan yang selama ini dikerjakan secara swadaya. Kegiatan ini juga "
            "menjadi ajang diskusi mengenai pengelolaan sampah dan kebersihan lingkungan di kawasan "
            "permukiman yang semakin padat, sekaligus penguatan semangat gotong royong sebagai "
            "kearifan lokal yang perlu terus dijaga."
        ),
    },
    {
        "no": 19, "hari_tanggal": "Selasa, 4 Agustus 2026", "waktu": "12.30-14.30 WIB",
        "tempat": "Desa Ngajum, Kecamatan Ngajum, Kabupaten Malang",
        "judul": "Dialog dan Kunjungan Lapangan di Desa Ngajum",
        "usulan": [
            "Mendesak perbaikan infrastruktur jalan desa yang menjadi akses utama distribusi hasil pertanian.",
            "Diusulkan peningkatan akses terhadap program bantuan pertanian seperti alat mesin pertanian dan bibit unggul.",
            "Berharap perhatian terhadap ketersediaan air bersih di musim kemarau bagi warga Desa Ngajum.",
        ],
        "fotos": ["R37.JPG", "R38.JPG"],
        "deskripsi_foto": (
            "Dialog dan kunjungan lapangan di Desa Ngajum, Kecamatan Ngajum, Kabupaten Malang, "
            "menghadirkan diskusi langsung dengan warga mengenai kondisi infrastruktur dan pertanian "
            "setempat. Warga menyampaikan bahwa jalan desa yang menjadi akses utama distribusi hasil "
            "pertanian masih dalam kondisi rusak dan perlu segera diperbaiki. Selain itu, warga "
            "berharap adanya bantuan alat mesin pertanian dan bibit unggul untuk meningkatkan hasil "
            "panen, serta perhatian khusus terhadap ketersediaan air bersih yang kerap menjadi "
            "masalah saat musim kemarau tiba."
        ),
    },
    {
        "no": 20, "hari_tanggal": "Selasa, 4 Agustus 2026", "waktu": "16.00-17.30 WIB",
        "tempat": "Kecamatan Lawang, Kabupaten Malang",
        "judul": "Menghadiri Festival Budaya Jaranan Kecamatan Lawang",
        "usulan": [
            "Berharap dukungan pemerintah daerah terhadap pelestarian seni budaya jaranan sebagai warisan budaya lokal.",
            "Diusulkan fasilitasi panggung dan anggaran penyelenggaraan festival budaya secara berkala guna menggerakkan ekonomi kreatif setempat.",
            "Meminta regenerasi pelaku seni jaranan melalui pelatihan bagi generasi muda agar tradisi ini tidak punah.",
        ],
        "fotos": ["R39.jpg", "R40.jpg"],
        "deskripsi_foto": (
            "Menghadiri Festival Budaya Jaranan di Kecamatan Lawang, Kabupaten Malang, menjadi "
            "penutup rangkaian kegiatan reses yang penuh dengan nuansa kebudayaan lokal. Para seniman "
            "dan pegiat budaya jaranan menyampaikan kekhawatiran atas menurunnya minat generasi muda "
            "untuk melestarikan kesenian tradisional ini akibat minimnya ruang pentas dan dukungan "
            "anggaran. Mereka berharap pemerintah daerah dapat memfasilitasi penyelenggaraan festival "
            "budaya secara berkala serta program regenerasi pelaku seni agar tradisi jaranan tetap "
            "lestari dan menjadi bagian dari penggerak ekonomi kreatif masyarakat Lawang."
        ),
    },
]

# ============================================================
# TEKS PENDAHULUAN (8 paragraf)
# ============================================================
PENDAHULUAN = [
    (
        "Sesuai dengan amanat konstitusi dan peraturan perundang-undangan, Anggota Dewan "
        "Perwakilan Rakyat Republik Indonesia memiliki kewajiban untuk melaksanakan kunjungan "
        "kerja reses ke daerah pemilihan pada setiap masa reses yang ditetapkan dalam kalender "
        "sidang DPR RI. Pada Masa Sidang V Tahun Sidang 2025-2026, masa reses dilaksanakan pada "
        "22 Juli hingga 13 Agustus 2026, dan dimanfaatkan untuk melaksanakan Kunjungan Kerja "
        "Reses Perseorangan di Daerah Pemilihan Jawa Timur V yang meliputi Kota Malang, Kabupaten "
        "Malang, dan Kota Batu. Kegiatan reses merupakan wujud nyata pelaksanaan fungsi "
        "representasi Anggota DPR RI dalam menjaga komunikasi dan kedekatan langsung dengan "
        "konstituen di daerah pemilihan."
    ),
    (
        "Salah satu fokus utama pelaksanaan reses kali ini adalah memantau dan mencari masukan "
        "langsung terkait kondisi perekonomian masyarakat di Malang Raya. Sebagai Anggota Komisi "
        "XI DPR RI yang membidangi keuangan dan perencanaan pembangunan, penting untuk melihat "
        "secara langsung bagaimana kondisi riil di lapangan, mulai dari daya beli masyarakat, "
        "keberlangsungan usaha koperasi dan UMKM, kinerja lembaga keuangan mikro seperti BPR dan "
        "koperasi simpan pinjam, hingga tantangan yang dihadapi sektor pertanian, perikanan, dan "
        "kehutanan yang menjadi tulang punggung perekonomian di sebagian besar wilayah Kabupaten "
        "Malang."
    ),
    (
        "Selain aspek ekonomi, reses ini juga diarahkan untuk memantau dan menggali kondisi sosial "
        "masyarakat secara lebih menyeluruh. Dinamika sosial yang berkembang di tengah masyarakat, "
        "baik pada aspek pendidikan, kehidupan keagamaan, kegotongroyongan, maupun pelestarian "
        "seni dan budaya lokal, menjadi bagian penting yang tidak dapat dipisahkan dari upaya "
        "membangun kesejahteraan masyarakat secara utuh. Kondisi sosial yang harmonis dan kohesif "
        "menjadi fondasi penting bagi keberlanjutan pembangunan ekonomi di suatu daerah."
    ),
    (
        "Melalui rangkaian dialog dan kunjungan lapangan, reses ini juga bertujuan untuk menggali "
        "secara langsung berbagai problematika yang dihadapi masyarakat sehari-hari. Persoalan "
        "seperti keterbatasan infrastruktur jalan dan irigasi, minimnya akses air bersih, "
        "keterbatasan sarana pendidikan dan kesehatan, hingga sulitnya akses permodalan bagi "
        "pelaku usaha kecil menjadi temuan yang konsisten muncul di berbagai titik kunjungan. "
        "Problematika tersebut perlu diidentifikasi secara mendalam agar dapat dirumuskan solusi "
        "kebijakan yang tepat sasaran."
    ),
    (
        "Momentum reses kali ini menjadi semakin strategis karena bersamaan dengan berlangsungnya "
        "pembahasan Rancangan Anggaran Pendapatan dan Belanja Negara (RAPBN) Tahun Anggaran 2027 "
        "di tingkat pusat. Sebagai Anggota Komisi XI yang turut membahas kebijakan fiskal dan "
        "anggaran negara, masukan-masukan yang diserap langsung dari masyarakat di daerah "
        "pemilihan menjadi bahan pengayaan yang sangat berharga untuk memastikan alokasi anggaran "
        "negara benar-benar responsif terhadap kebutuhan riil masyarakat di daerah."
    ),
    (
        "Laporan ini secara khusus memuat rangkaian kegiatan reses yang dilaksanakan pada periode "
        "27 Juli hingga 4 Agustus 2026, terdiri atas 20 kegiatan yang tersebar di wilayah Kota "
        "Malang, Kabupaten Malang, dan Kota Batu. Rangkaian kegiatan tersebut mencakup berbagai "
        "elemen masyarakat, mulai dari pelaku ekonomi (koperasi, UMKM, BPR, industri jasa "
        "keuangan), kalangan pendidikan, tokoh dan umat lintas agama, jajaran struktur partai, "
        "kelompok tani dan nelayan, hingga pegiat seni dan budaya lokal."
    ),
    (
        "Pelaksanaan reses dilakukan melalui berbagai metode, antara lain dialog interaktif, "
        "kunjungan lapangan, serta pertemuan formal dengan mitra kerja dan elemen masyarakat di "
        "daerah pemilihan. Pendekatan yang partisipatif dan terbuka ini dimaksudkan agar aspirasi "
        "yang terserap benar-benar mencerminkan kebutuhan riil masyarakat, tanpa tersaring oleh "
        "kepentingan pihak tertentu, sehingga dapat menjadi bahan yang otentik dalam perumusan "
        "kebijakan di tingkat nasional."
    ),
    (
        "Laporan ini disusun sebagai bentuk pertanggungjawaban pelaksanaan tugas konstitusional "
        "sekaligus dokumentasi menyeluruh atas seluruh rangkaian kegiatan Kunjungan Kerja Reses "
        "Perseorangan yang telah dilaksanakan. Melalui laporan ini, diharapkan seluruh aspirasi, "
        "masukan, dan usulan yang telah diserap dari masyarakat Daerah Pemilihan Jawa Timur V "
        "dapat ditindaklanjuti secara konkret, baik melalui forum rapat kerja Komisi XI, "
        "pembahasan RAPBN 2027, maupun koordinasi dengan instansi pemerintah terkait."
    ),
]

# ============================================================
# TEKS PENUTUP (6 paragraf)
# ============================================================
PENUTUP = [
    (
        "Rangkaian Kunjungan Kerja Reses Perseorangan yang dilaksanakan pada 27 Juli hingga 4 "
        "Agustus 2026, dalam Masa Sidang V Tahun Sidang 2025-2026, telah berhasil menjangkau 20 "
        "kegiatan yang tersebar di Kota Malang, Kabupaten Malang, dan Kota Batu. Pelaksanaan reses "
        "ini berhasil mencapai tujuan awal, yakni memantau dan mencari masukan terkait kondisi "
        "perekonomian masyarakat, menggali kondisi sosial, mengidentifikasi problematika yang "
        "dihadapi masyarakat, serta menghimpun masukan sebagai bahan pengayaan pembahasan RAPBN "
        "Tahun Anggaran 2027."
    ),
    (
        "Dari sisi perekonomian, dialog dengan pengurus koperasi, pelaku UMKM, industri jasa "
        "keuangan, hingga Bank Perkreditan Rakyat menunjukkan bahwa akses permodalan dengan syarat "
        "yang mudah dan bunga terjangkau masih menjadi kebutuhan sentral bagi pelaku usaha di "
        "Malang Raya. Persoalan literasi keuangan, ancaman pinjaman ilegal, serta perlunya "
        "regulasi yang lebih berpihak pada lembaga keuangan mikro menjadi catatan penting yang "
        "harus dikawal dalam pembahasan kebijakan fiskal ke depan."
    ),
    (
        "Dari sisi sosial, pendidikan, dan keagamaan, dialog dengan orang tua dan siswa SMA/SMK, "
        "umat Katolik di Paroki Kota Batu, jemaat gereja di Kota Malang, hingga warga yang "
        "melaksanakan kerja bakti serta pegiat seni jaranan di Kecamatan Lawang menunjukkan bahwa "
        "kohesi sosial dan pelestarian nilai budaya lokal masih menjadi kekuatan utama masyarakat "
        "Malang Raya. Kebutuhan akan dukungan sarana pendidikan, fasilitas keagamaan, dan ruang "
        "berkesenian perlu terus diperjuangkan agar modal sosial ini tidak tergerus."
    ),
    (
        "Sementara itu, dari sisi infrastruktur dan sektor primer, temuan di sejumlah desa seperti "
        "Druju, Sambigede, Kalipare, dan Ngajum, serta dialog dengan petani dan nelayan di "
        "Sumbermanjing Wetan dan kunjungan ke Perum Perhutani KPH Malang, menegaskan bahwa "
        "pemerataan infrastruktur dasar seperti jalan, irigasi, air bersih, dan sarana produksi "
        "pertanian maupun perikanan masih menjadi kebutuhan mendesak yang harus menjadi prioritas "
        "alokasi anggaran pembangunan daerah."
    ),
    (
        "Seluruh aspirasi dan usulan yang telah dihimpun selama pelaksanaan reses ini akan menjadi "
        "bahan pengayaan dalam pembahasan RAPBN Tahun Anggaran 2027 serta agenda kerja Komisi XI "
        "DPR RI, khususnya yang berkaitan dengan kebijakan fiskal, sektor keuangan, dan "
        "perencanaan pembangunan nasional. Aspirasi yang bersifat teknis dan kewenangan daerah "
        "juga akan dikoordinasikan dengan pemerintah daerah dan instansi terkait agar dapat "
        "ditindaklanjuti secara konkret dan tepat sasaran."
    ),
    (
        "Demikian laporan Kunjungan Kerja Reses Perseorangan ini disusun sebagai bentuk "
        "pertanggungjawaban pelaksanaan tugas konstitusional selama Masa Reses Masa Sidang V "
        "Tahun Sidang 2025-2026. Ucapan terima kasih disampaikan kepada seluruh elemen masyarakat, "
        "tokoh agama, pengurus koperasi dan lembaga keuangan, kalangan pendidikan, jajaran partai, "
        "petani, nelayan, serta pegiat seni budaya di Daerah Pemilihan Jawa Timur V atas "
        "partisipasi dan keterbukaannya dalam menyampaikan aspirasi. Komitmen untuk terus "
        "memperjuangkan aspirasi masyarakat akan senantiasa dijalankan dengan sungguh-sungguh "
        "dalam setiap forum kedewanan di DPR RI."
    ),
]

# ============================================================
# DASAR HUKUM
# ============================================================
DASAR_HUKUM = [
    "Pasal 20A ayat (3): Memberikan hak kepada setiap Anggota DPR untuk mengajukan pertanyaan, menyampaikan usul dan pendapat, serta hak imunitas dalam menjalankan tugas kedewanan, termasuk dalam masa reses.",
]
DASAR_HUKUM_MD3 = [
    "Anggota DPR RI mempunyai hak dan kewajiban untuk menyerap, menghimpun, menampung, dan menindaklanjuti aspirasi masyarakat, salah satunya melalui pelaksanaan kunjungan kerja reses ke daerah pemilihan.",
]
DASAR_HUKUM_TATIB = [
    "Mengatur bahwa masa reses merupakan bagian dari masa persidangan yang digunakan oleh Anggota DPR RI untuk melaksanakan kunjungan kerja perseorangan ke daerah pemilihan guna menyerap, menghimpun, dan menindaklanjuti aspirasi masyarakat.",
]

# ============================================================
# TUJUAN
# ============================================================
TUJUAN = [
    "Menyerap Aspirasi Masyarakat: Menghimpun secara langsung aspirasi, keluhan, dan usulan masyarakat di Daerah Pemilihan Jawa Timur V terkait kondisi ekonomi, sosial, dan pembangunan daerah.",
    "Memantau Kondisi Perekonomian: Melihat secara langsung dampak dinamika makroekonomi terhadap pelaku usaha, koperasi, lembaga keuangan mikro, serta sektor pertanian, perikanan, dan kehutanan di tingkat lokal.",
    "Menggali Kondisi Sosial: Memantau dinamika sosial, pendidikan, kehidupan keagamaan, dan pelestarian seni budaya sebagai bagian dari upaya menjaga kohesi sosial masyarakat.",
    "Pengayaan Bahan Pembahasan RAPBN 2027: Menghimpun masukan masyarakat sebagai bahan pengayaan pembahasan Rancangan Anggaran Pendapatan dan Belanja Negara Tahun Anggaran 2027 di Komisi XI DPR RI.",
    "Menindaklanjuti Aspirasi: Menjadikan seluruh hasil serapan aspirasi sebagai dasar advokasi kebijakan dalam forum rapat kerja, rapat dengar pendapat, maupun sidang paripurna di DPR RI.",
]

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                            bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=0, space_after=6, color=None, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_photo_to_doc(doc, foto_path, width_cm=13):
    if not os.path.exists(foto_path):
        print(f"  [SKIP] Foto tidak ditemukan: {foto_path}")
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(foto_path, width=Cm(width_cm))
    except Exception as e:
        print(f"  [ERROR] Gagal menambahkan foto {foto_path}: {e}")


def set_table_cell_lines(cell, lines, font_name="Times New Roman", font_size=10,
                         bold=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Isi cell tabel dengan satu atau lebih baris/paragraf."""
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
            rPr.append(rFonts)
        else:
            rFonts.set(qn('w:eastAsia'), font_name)


# ============================================================
# MAIN DOCUMENT CREATION
# ============================================================

def create_document():
    doc = Document()

    # -- Page setup: A4, Portrait --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========================================
    # SAMPUL / COVER
    # ========================================
    for _ in range(3):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "DEWAN PERWAKILAN RAKYAT REPUBLIK INDONESIA",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    add_formatted_paragraph(
        doc, "LAPORAN KUNJUNGAN KERJA RESES PERSEORANGAN",
        font_size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    add_formatted_paragraph(
        doc, MASA_SIDANG,
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
    )

    for _ in range(4):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "Nama                        : Ir. Andreas Eddy Susetyo, M.M.",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, "Nomor Anggota      : A-220",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, "Komisi                       : XI",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, "Fraksi                        : PDI Perjuangan",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, f"Daerah Pemilihan     : {DAPIL}",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4
    )
    add_formatted_paragraph(
        doc, f"Masa Reses             : {PERIODE_RESES}",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
    )

    for _ in range(6):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "DEWAN PERWAKILAN RAKYAT REPUBLIK INDONESIA TAHUN 2026",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0
    )

    doc.add_page_break()

    # ========================================
    # A. PENDAHULUAN
    # ========================================
    add_formatted_paragraph(
        doc, "A. PENDAHULUAN",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6
    )

    for para_text in PENDAHULUAN:
        add_formatted_paragraph(
            doc, para_text,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

    # ========================================
    # B. DASAR HUKUM
    # ========================================
    add_formatted_paragraph(
        doc, "B. DASAR HUKUM",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc, "Pelaksanaan Kunjungan Kerja Reses Perseorangan ini berlandaskan pada peraturan "
        "perundang-undangan sebagai berikut:",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6
    )

    add_formatted_paragraph(
        doc, "Undang-Undang Dasar Negara Republik Indonesia Tahun 1945 (UUD 1945)",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4
    )
    for dh in DASAR_HUKUM:
        add_formatted_paragraph(
            doc, dh, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    add_formatted_paragraph(
        doc, "Undang-Undang Nomor 17 Tahun 2014 tentang MPR, DPR, DPD, dan DPRD (UU MD3)",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=8, space_after=4
    )
    for dh in DASAR_HUKUM_MD3:
        add_formatted_paragraph(
            doc, dh, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    add_formatted_paragraph(
        doc, "Peraturan DPR RI Nomor 1 Tahun 2025 tentang Tata Tertib",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=8, space_after=4
    )
    for dh in DASAR_HUKUM_TATIB:
        add_formatted_paragraph(
            doc, dh, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    # ========================================
    # C. TUJUAN
    # ========================================
    add_formatted_paragraph(
        doc, "C. TUJUAN",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc, "Pelaksanaan Kunjungan Kerja Reses Perseorangan ini memiliki tujuan sebagai berikut:",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4
    )

    for t in TUJUAN:
        add_formatted_paragraph(
            doc, t, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    # ========================================
    # D. PELAKSANAAN KUNJUNGAN KERJA
    # ========================================
    doc.add_page_break()
    add_formatted_paragraph(
        doc, "D. PELAKSANAAN KUNJUNGAN KERJA",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6
    )

    add_formatted_paragraph(
        doc,
        f"Kunjungan Kerja Reses Perseorangan dilaksanakan pada periode {PERIODE_RESES}, dengan "
        f"rangkaian kegiatan pada tanggal 27 Juli hingga 4 Agustus 2026 sebagaimana tercantum "
        f"dalam tabel berikut:",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=8, first_line_indent=1.27
    )

    num_rows = len(KEGIATAN) + 1
    table = doc.add_table(rows=num_rows, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Cm(1.0), Cm(3.3), Cm(3.86), Cm(6.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    headers = ["NO", "KEGIATAN RESES", "WAKTU & TEMPAT", "KETERANGAN/USULAN"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9E2F3")
        set_table_cell_lines(cell, [header], bold=True, font_size=10,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for i, keg in enumerate(KEGIATAN):
        row = table.rows[i + 1]
        set_table_cell_lines(row.cells[0], [str(keg["no"])], font_size=10,
                             alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_table_cell_lines(row.cells[1], [keg["judul"]], font_size=10)
        set_table_cell_lines(
            row.cells[2],
            [f'{keg["hari_tanggal"]}', f'Pukul: {keg["waktu"]}', f'Tempat: {keg["tempat"]}'],
            font_size=10
        )
        usulan_lines = [f"{j + 1}. {u}" for j, u in enumerate(keg["usulan"])]
        set_table_cell_lines(row.cells[3], usulan_lines, font_size=10)

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "000000"},
                bottom={"val": "single", "sz": "4", "color": "000000"},
                start={"val": "single", "sz": "4", "color": "000000"},
                end={"val": "single", "sz": "4", "color": "000000"},
            )

    # ========================================
    # E. PENUTUP
    # ========================================
    doc.add_page_break()
    add_formatted_paragraph(
        doc, "E. PENUTUP",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6
    )

    for para_text in PENUTUP:
        add_formatted_paragraph(
            doc, para_text,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

    # ========================================
    # TANDA TANGAN
    # ========================================
    add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, TANGGAL_LAPORAN,
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    for _ in range(3):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "Ir. Andreas Eddy Susetyo, MM",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    add_formatted_paragraph(
        doc, "Anggota DPR RI A–220",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    # ========================================
    # DOKUMENTASI FOTO KEGIATAN
    # ========================================
    doc.add_page_break()
    add_formatted_paragraph(
        doc, "DOKUMENTASI FOTO KEGIATAN",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18
    )

    for keg in KEGIATAN:
        add_formatted_paragraph(
            doc, f'Kegiatan {keg["no"]}: {keg["judul"]}',
            font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=12, space_after=2
        )
        add_formatted_paragraph(
            doc, f'{keg["hari_tanggal"]}, {keg["waktu"]} - {keg["tempat"]}',
            font_size=11, italic=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_after=6
        )

        add_formatted_paragraph(
            doc, keg["deskripsi_foto"],
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

        for foto_name in keg["fotos"]:
            foto_path = os.path.join(FOTO_DIR, foto_name)
            add_photo_to_doc(doc, foto_path, width_cm=13)

        add_formatted_paragraph(doc, "", font_size=6, space_after=0)

    # ========================================
    # SAVE
    # ========================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Laporan berhasil dibuat: {OUTPUT_FILE}")
    print(f"     Jumlah kegiatan: {len(KEGIATAN)}")
    print(f"     Jumlah foto: {sum(len(k['fotos']) for k in KEGIATAN)}")


if __name__ == "__main__":
    create_document()
