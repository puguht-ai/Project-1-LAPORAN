"""
Script untuk membuat Laporan Kegiatan Rumah Aspirasi DPR RI
Bulan Juni 2026 - Ir. Andreas Eddy Susetyo, MM (A-220)
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from PIL import Image

# ============================================================
# CONFIG
# ============================================================
OUTPUT_DIR = r"D:\BLACKBOX DATA"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Laporan Rumah Aspirasi Juni 2026.docx")
FOTO_DIR = r"D:\DATA PAPA\LAPORAN\RUMAH ASPIRASI\JUNI 2026"

# ============================================================
# DATA KEGIATAN
# ============================================================
KEGIATAN = [
    {
        "no": 1,
        "tanggal": "3 Juni 2026",
        "kegiatan": "Rapat Evaluasi dan Penyusunan Rencana Program Bulan Juni 2026",
        "tempat": "Rumah Aspirasi, Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang",
        "keterangan": "Melaksanakan rapat evaluasi kinerja semester pertama dan menyusun rencana program kerja bulan Juni 2026 untuk memastikan seluruh agenda Rumah Aspirasi berjalan efektif dan tepat sasaran.",
        "fotos": ["A1.jpeg", "A2.jpeg"],
        "deskripsi_foto": (
            "Rapat Evaluasi dan Penyusunan Rencana Program Bulan Juni 2026 di Rumah Aspirasi, "
            "Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang. "
            "Rapat ini menjadi momentum penting untuk mengevaluasi capaian kinerja Rumah Aspirasi "
            "selama semester pertama tahun 2026. Tim melakukan review menyeluruh terhadap seluruh "
            "program yang telah dilaksanakan, mengidentifikasi keberhasilan maupun kendala yang "
            "dihadapi di lapangan. Pembahasan juga mencakup pemetaan kebutuhan masyarakat terkini "
            "di wilayah Malang Raya yang perlu mendapat perhatian khusus. Hasil evaluasi ini menjadi "
            "dasar penyusunan rencana program bulan Juni yang lebih terarah dan responsif terhadap "
            "dinamika aspirasi konstituen."
        ),
    },
    {
        "no": 2,
        "tanggal": "9 Juni 2026",
        "kegiatan": "Menerima Kunjungan dari Pengurus Gereja Wilayah Kepanjen",
        "tempat": "Rumah Aspirasi, Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang",
        "keterangan": "Menerima kunjungan silaturahmi dari pengurus Gereja wilayah Kepanjen, Kab. Malang untuk berdiskusi mengenai aspirasi umat dan kebutuhan pembangunan fasilitas keagamaan.",
        "fotos": ["A3.jpeg", "A4.jpeg"],
        "deskripsi_foto": (
            "Menerima Kunjungan dari Pengurus Gereja Wilayah Kepanjen, Kab. Malang di Rumah Aspirasi, "
            "Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang. "
            "Kunjungan ini merupakan wujud keterbukaan Rumah Aspirasi dalam menerima aspirasi dari "
            "seluruh elemen masyarakat tanpa memandang latar belakang. Pengurus Gereja menyampaikan "
            "berbagai harapan dan kebutuhan umat di wilayah Kepanjen, termasuk perhatian terhadap "
            "pembangunan sarana ibadah dan kegiatan sosial kemasyarakatan. Dialog berlangsung hangat "
            "dan konstruktif, mencerminkan semangat kerukunan antarumat beragama yang menjadi ciri "
            "khas masyarakat Malang Raya. Pertemuan ini juga memperkuat komitmen untuk menjaga "
            "keharmonisan dan toleransi dalam kehidupan bermasyarakat."
        ),
    },
    {
        "no": 3,
        "tanggal": "17 Juni 2026",
        "kegiatan": "Menerima Kunjungan Komunitas Rider Malang Raya",
        "tempat": "Rumah Aspirasi, Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang",
        "keterangan": "Menerima kunjungan silaturahmi dari Komunitas Rider Malang Raya untuk berdiskusi mengenai aspirasi komunitas terkait keselamatan berkendara dan pengembangan wisata otomotif daerah.",
        "fotos": ["A5.jpeg", "A6.jpeg"],
        "deskripsi_foto": (
            "Menerima Kunjungan Komunitas Rider Malang Raya di Rumah Aspirasi, "
            "Dusun Krajan, Desa Karang Tengah, Kec. Pakis, Kab. Malang. "
            "Komunitas Rider Malang Raya yang terdiri dari para penggemar sepeda motor dari berbagai "
            "klub di wilayah Malang Raya hadir untuk menyampaikan aspirasi dan berdiskusi mengenai "
            "berbagai isu yang berkaitan dengan komunitas mereka. Pembahasan meliputi keselamatan "
            "berkendara, kondisi infrastruktur jalan, serta potensi pengembangan wisata otomotif "
            "yang dapat mendukung perekonomian daerah. Pertemuan ini menunjukkan bahwa Rumah Aspirasi "
            "terbuka bagi seluruh kalangan masyarakat, termasuk komunitas hobi yang memiliki kontribusi "
            "positif bagi kehidupan sosial dan ekonomi masyarakat Malang Raya."
        ),
    },
    {
        "no": 4,
        "tanggal": "21 Juni 2026",
        "kegiatan": "Menghadiri Acara Komunitas UMKM Kota Malang",
        "tempat": "Restoran Tugu, Kota Malang",
        "keterangan": "Menghadiri undangan acara silaturahmi dan diskusi bersama Komunitas UMKM Kota Malang untuk mendengarkan aspirasi pelaku usaha mikro, kecil, dan menengah terkait pengembangan usaha dan akses permodalan.",
        "fotos": ["A7.jpg", "A8.jpg"],
        "deskripsi_foto": (
            "Menghadiri Acara Komunitas UMKM Kota Malang di Restoran Tugu, Kota Malang. "
            "Kehadiran dalam acara ini merupakan bentuk dukungan nyata terhadap perkembangan sektor "
            "UMKM yang menjadi tulang punggung perekonomian masyarakat Kota Malang. Para pelaku UMKM "
            "menyampaikan berbagai aspirasi terkait kemudahan akses permodalan, pelatihan digital "
            "marketing, serta perlindungan hukum bagi usaha kecil. Diskusi berlangsung interaktif "
            "dengan berbagai masukan konstruktif yang akan ditindaklanjuti melalui mekanisme "
            "kelembagaan di DPR RI. Acara ini juga menjadi ajang networking antarpeaku UMKM "
            "untuk saling berbagi pengalaman dan memperluas jaringan bisnis mereka."
        ),
    },
    {
        "no": 5,
        "tanggal": "28 Juni 2026",
        "kegiatan": "Menghadiri Acara Silaturahmi Warga Tionghoa Malang Raya",
        "tempat": "Hotel Ijen Suite, Kota Malang",
        "keterangan": "Tim Rumah Aspirasi diundang menghadiri acara silaturahmi warga Tionghoa Malang Raya sebagai wujud kebersamaan dan kerukunan antaretnis dalam membangun Malang Raya yang inklusif.",
        "fotos": ["a9.JPG", "A10.JPG"],
        "deskripsi_foto": (
            "Menghadiri Acara Silaturahmi Warga Tionghoa Malang Raya di Hotel Ijen Suite, Kota Malang. "
            "Kehadiran Tim Rumah Aspirasi dalam acara ini mencerminkan komitmen untuk merangkul "
            "seluruh elemen masyarakat tanpa membedakan suku, agama, ras, dan antargolongan. "
            "Warga Tionghoa Malang Raya menyampaikan berbagai harapan terkait iklim usaha yang "
            "kondusif, kerukunan antaretnis, serta kontribusi mereka dalam pembangunan daerah. "
            "Acara silaturahmi ini berlangsung dalam suasana yang hangat dan penuh persaudaraan, "
            "menunjukkan bahwa keberagaman adalah kekuatan yang harus terus dijaga dan dikembangkan. "
            "Pertemuan ini juga memperkuat jalinan komunikasi antara wakil rakyat dengan seluruh "
            "komponen masyarakat di Malang Raya."
        ),
    },
]

# ============================================================
# TEKS PENDAHULUAN (~3.500 karakter)
# ============================================================
PENDAHULUAN = [
    (
        "Sesuai wewenang Dewan Perwakilan Rakyat Republik Indonesia yang termaktub pada "
        "UU MD3 Pasal 72, Pasal 81, Pasal 234 atau secara detail dalam Tata Tertib DPR "
        "dalam Pasal 1 yang menyebutkan bahwa rumah aspirasi adalah kantor setiap anggota "
        "sebagai tempat penyerapan aspirasi konstituen yang berada di daerah pemilihan anggota. "
        "Melalui wadah ini, Anggota Dewan diharapkan dapat berkomunikasi secara langsung dan "
        "intens dengan konstituennya guna menjalankan fungsi wakil rakyat dalam menampung "
        "aspirasi masyarakat."
    ),
    (
        "Memasuki bulan Juni 2026, Rumah Aspirasi memasuki fase penting yakni evaluasi "
        "paruh tahun pertama masa kerja. Enam bulan perjalanan telah memberikan banyak "
        "pelajaran berharga mengenai dinamika aspirasi masyarakat di Daerah Pemilihan "
        "Jawa Timur V. Konsolidasi internal menjadi prioritas utama untuk memastikan bahwa "
        "seluruh program yang telah direncanakan berjalan sesuai target dan memberikan "
        "dampak nyata bagi konstituen. Evaluasi ini bukan sekadar formalitas administratif, "
        "melainkan refleksi mendalam atas efektivitas kinerja dalam menyerap dan "
        "menindaklanjuti aspirasi rakyat."
    ),
    (
        "Selain kegiatan konsolidasi internal, bulan Juni juga diwarnai dengan berbagai "
        "kunjungan dan undangan dari beragam elemen masyarakat. Kehadiran pengurus Gereja "
        "dari wilayah Kepanjen menunjukkan bahwa Rumah Aspirasi telah menjadi ruang yang "
        "inklusif bagi seluruh kalangan tanpa memandang latar belakang agama maupun etnis. "
        "Demikian pula kunjungan dari Komunitas Rider Malang Raya membuktikan bahwa aspirasi "
        "masyarakat tidak hanya datang dari kelompok formal, tetapi juga dari komunitas-komunitas "
        "yang memiliki kepedulian terhadap pembangunan daerah."
    ),
    (
        "Undangan dari Komunitas UMKM Kota Malang dan acara silaturahmi warga Tionghoa "
        "Malang Raya menjadi bukti bahwa jangkauan Rumah Aspirasi semakin meluas ke berbagai "
        "sektor kehidupan masyarakat. Sektor UMKM yang menjadi tulang punggung perekonomian "
        "lokal memerlukan perhatian serius dari wakil rakyat, terutama dalam hal akses "
        "permodalan, pelatihan, dan perlindungan usaha. Sementara itu, kehadiran dalam acara "
        "warga Tionghoa menegaskan komitmen untuk membangun Malang Raya yang inklusif dan "
        "menghargai keberagaman sebagai kekuatan bersama."
    ),
    (
        "Pengelolaan Rumah Aspirasi juga berkaitan erat dengan prinsip akuntabilitas dan "
        "transparansi penggunaan fasilitas negara yang telah dipercayakan kepada anggota "
        "dewan. Setiap kegiatan yang dilaksanakan harus dapat dipertanggungjawabkan secara "
        "administratif maupun substansif kepada publik dan lembaga DPR RI. Hal ini menjadi "
        "standar operasional yang wajib dipenuhi untuk menjaga integritas lembaga perwakilan "
        "rakyat di mata masyarakat."
    ),
    (
        "Laporan kegiatan bulan Juni 2026 ini disusun sebagai bukti fisik dan naratif atas "
        "seluruh rangkaian kegiatan yang telah dilaksanakan selama satu bulan periode berjalan. "
        "Melalui dokumen ini, kami berharap dapat memberikan gambaran yang utuh mengenai kinerja "
        "penyerapan aspirasi serta komitmen kami dalam melayani konstituen di wilayah Malang Raya. "
        "Besar harapan kami laporan ini dapat menjadi sarana evaluasi dan peningkatan kualitas "
        "pelayanan publik di masa mendatang."
    ),
]

# ============================================================
# TEKS PENUTUP (~3.000 karakter)
# ============================================================
PENUTUP = [
    (
        "Demikian laporan kegiatan Rumah Aspirasi ini disusun sebagai wujud nyata akuntabilitas "
        "dan transparansi kinerja Ir. Andreas Eddy Susetyo, MM dalam menjalankan amanah sebagai "
        "wakil rakyat selama bulan Juni 2026. Laporan ini merupakan bagian dari kewajiban "
        "konstitusional untuk melaporkan setiap upaya penyerapan aspirasi yang telah dilaksanakan "
        "di Daerah Pemilihan Jawa Timur V."
    ),
    (
        "Bulan Juni menjadi titik refleksi penting setelah enam bulan perjalanan Rumah Aspirasi "
        "di tahun 2026. Evaluasi internal yang dilakukan menunjukkan bahwa program-program yang "
        "telah dijalankan mendapat respons positif dari masyarakat. Berbagai kunjungan dari "
        "beragam elemen masyarakat — mulai dari pengurus Gereja, Komunitas Rider, pelaku UMKM, "
        "hingga warga Tionghoa — membuktikan bahwa Rumah Aspirasi telah berhasil memposisikan "
        "diri sebagai ruang publik yang terbuka, inklusif, dan dipercaya oleh konstituen."
    ),
    (
        "Keberagaman latar belakang tamu dan aspirasi yang diterima selama bulan Juni menjadi "
        "modal sosial yang sangat berharga. Aspirasi dari komunitas UMKM terkait akses permodalan "
        "dan pelatihan digital akan ditindaklanjuti melalui koordinasi dengan instansi terkait "
        "dan pembahasan di Komisi XI DPR RI. Sementara itu, harapan dari berbagai komunitas "
        "keagamaan dan etnis untuk menjaga kerukunan dan toleransi akan terus menjadi perhatian "
        "utama dalam setiap kebijakan yang diperjuangkan."
    ),
    (
        "Rumah Aspirasi akan terus konsisten membuka pintu bagi seluruh lapisan masyarakat guna "
        "memperkuat partisipasi publik dalam proses demokrasi yang sehat. Memasuki semester kedua "
        "tahun 2026, kami bertekad untuk meningkatkan intensitas dan kualitas penyerapan aspirasi "
        "dengan memperluas jangkauan ke wilayah-wilayah yang belum terjangkau secara optimal. "
        "Semangat kebersamaan dalam keberagaman yang tercermin dari kegiatan bulan Juni ini akan "
        "menjadi landasan kuat untuk terus melayani rakyat dengan sepenuh hati."
    ),
    (
        "Akhir kata, kami menyampaikan apresiasi kepada semua pihak yang telah berpartisipasi "
        "dan membantu kelancaran rangkaian kegiatan ini. Semoga upaya yang dilakukan melalui "
        "Rumah Aspirasi ini senantiasa memberikan manfaat yang nyata bagi masyarakat dan terus "
        "menjadi sarana perjuangan bagi aspirasi rakyat yang lebih baik di masa depan."
    ),
]

# ============================================================
# HELPER FUNCTIONS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, font_name="Times New Roman", font_size=12,
                            bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                            space_before=0, space_after=6, color=None, first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)
    return p


def add_photo_to_doc(doc, foto_path, width_cm=14):
    """Add a photo centered in the document."""
    if not os.path.exists(foto_path):
        print(f"  [SKIP] Foto tidak ditemukan: {foto_path}")
        return
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(foto_path, width=Cm(width_cm))
    except Exception as e:
        print(f"  [ERROR] Gagal menambahkan foto {foto_path}: {e}")


def set_table_cell_text(cell, text, font_name="Times New Roman", font_size=11,
                        bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    """Set text in a table cell with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), font_name)


# ============================================================
# MAIN DOCUMENT CREATION
# ============================================================

def create_document():
    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # ========================================
    # SAMPUL / COVER
    # ========================================
    for _ in range(4):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "DEWAN PERWAKILAN RAKYAT REPUBLIK INDONESIA",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12
    )

    add_formatted_paragraph(doc, "", font_size=14, space_after=0)

    add_formatted_paragraph(
        doc, "LAPORAN KEGIATAN RUMAH ASPIRASI DPR RI",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    add_formatted_paragraph(
        doc, "BULAN JUNI TAHUN 2026 TAHUN ANGGARAN 2026",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
    )

    for _ in range(6):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "IR. ANDREAS EDDY SUSETYO, MM",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6
    )

    add_formatted_paragraph(
        doc, "ANGGOTA FRAKSI PDI PERJUANGAN A-220",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24
    )

    for _ in range(6):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "DEWAN PERWAKILAN RAKYAT REPUBLIK INDONESIA TAHUN 2026",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0
    )

    # Page break after cover
    doc.add_page_break()

    # ========================================
    # A. PENDAHULUAN
    # ========================================
    add_formatted_paragraph(
        doc, "A. PENDAHULUAN",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6
    )

    for para_text in PENDAHULUAN:
        add_formatted_paragraph(
            doc, para_text,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

    # ========================================
    # B. DASAR HUKUM
    # ========================================
    add_formatted_paragraph(
        doc, "B. DASAR HUKUM",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc, "Peraturan DPR RI No 1 Tahun 2020:",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=4
    )

    dasar_hukum = [
        "Pasal 1 ayat(17): Rumah Aspirasi adalah kantor setiap Anggota sebagai tempat penyerapan aspirasi rakyat di daerah pemilihan.",
        "Pasal 238 ayat(4): Anggota dapat membuat rumah aspirasi untuk membuka ruang partisipasi publik.",
        "Pasal 241 ayat(1): Anggota dibantu oleh Tenaga Ahli dan staf administrasi dalam mengelola rumah aspirasi.",
        "Pasal 241 ayat(2)&(3): Rumah Aspirasi didukung oleh anggaran DPR yang dikelola sesuai peraturan pertanggungjawaban pengelolaan anggaran.",
    ]
    for dh in dasar_hukum:
        add_formatted_paragraph(
            doc, dh,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    # ========================================
    # C. TUJUAN
    # ========================================
    add_formatted_paragraph(
        doc, "C. TUJUAN",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc, "Kegiatan Rumah Aspirasi ini memiliki tujuan strategis sebagai berikut:",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4
    )

    tujuan = [
        "Optimalisasi Representasi: Menjadi jembatan komunikasi dua arah yang efektif untuk menyerap, menghimpun, dan menindaklanjuti aspirasi riil masyarakat di Daerah Pemilihan Jawa Timur V (Malang Raya) secara sistematis.",
        "Pemberdayaan Komunitas/Masyarakat Lokal: Membina dan mendampingi berbagai komunitas, baik komunitas pemuda, kelompok tani, maupun elemen masyarakat lainnya agar mampu berpartisipasi aktif dalam pembangunan daerah.",
        "Advokasi Kebijakan: Menjadikan masukan dari masyarakat sebagai bahan kajian mendalam untuk diperjuangkan dalam rapat-rapat komisi maupun sidang paripurna di DPR RI guna menghasilkan kebijakan yang pro-rakyat.",
        "Pusat Informasi dan Solusi: Menyediakan ruang publik bagi masyarakat untuk mendapatkan informasi terkait program pemerintah serta mencari solusi atas permasalahan sosial-ekonomi yang dihadapi di tingkat lokal.",
    ]
    for t in tujuan:
        add_formatted_paragraph(
            doc, t,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=4, first_line_indent=1.27
        )

    # ========================================
    # D. KEGIATAN
    # ========================================
    add_formatted_paragraph(
        doc, "D. KEGIATAN",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc,
        "Kegiatan ini dijalankan untuk menampung aspirasi masyarakat yang merupakan amanat "
        "konstitusi sebagai wakil rakyat menyangkut berbagai aspek kebutuhan mendasar masyarakat, "
        "terutama yang berkaitan dengan tugas kedewanan yakni fungsi pengawasan, anggaran dan legislasi.",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6, first_line_indent=1.27
    )

    # ========================================
    # E. TEMPAT DAN WAKTU
    # ========================================
    add_formatted_paragraph(
        doc, "E. TEMPAT DAN WAKTU",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    add_formatted_paragraph(
        doc,
        "Kegiatan dilaksanakan di wilayah Malang Raya (Kabupaten Malang, Kota Malang, Kota Batu) "
        "untuk memastikan seluruh aspirasi tersalurkan kepada pihak berkompeten agar dapat "
        "terealisasi sesuai kebutuhan di lapangan.",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        space_after=6, first_line_indent=1.27
    )

    # ========================================
    # F. KEGIATAN RUMAH ASPIRASI BULAN JUNI 2026
    # ========================================
    add_formatted_paragraph(
        doc, "F. KEGIATAN RUMAH ASPIRASI BULAN JUNI 2026",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=12, space_after=6
    )

    # Create table
    num_rows = len(KEGIATAN) + 1  # +1 for header
    table = doc.add_table(rows=num_rows, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    col_widths = [Cm(1.2), Cm(2.5), Cm(4.5), Cm(4.0), Cm(4.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Header row
    headers = ["NO", "TANGGAL", "KEGIATAN", "TEMPAT", "KETERANGAN"]
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, "D9E2F3")
        set_table_cell_text(cell, header, bold=True, font_size=10)

    # Data rows
    for i, keg in enumerate(KEGIATAN):
        row = table.rows[i + 1]
        set_table_cell_text(row.cells[0], str(keg["no"]), font_size=10)
        set_table_cell_text(row.cells[1], keg["tanggal"], font_size=10)
        set_table_cell_text(row.cells[2], keg["kegiatan"], font_size=10,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[3], keg["tempat"], font_size=10,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)
        set_table_cell_text(row.cells[4], keg["keterangan"], font_size=10,
                           alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # Set borders for all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "000000"},
                bottom={"val": "single", "sz": "4", "color": "000000"},
                start={"val": "single", "sz": "4", "color": "000000"},
                end={"val": "single", "sz": "4", "color": "000000"},
            )

    # ========================================
    # G. PENUTUP
    # ========================================
    add_formatted_paragraph(
        doc, "G. PENUTUP",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
        space_before=18, space_after=6
    )

    for para_text in PENUTUP:
        add_formatted_paragraph(
            doc, para_text,
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

    # ========================================
    # TANDA TANGAN
    # ========================================
    add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "Malang, 30 Juni 2026",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    for _ in range(3):
        add_formatted_paragraph(doc, "", font_size=12, space_after=0)

    add_formatted_paragraph(
        doc, "Ir. Andreas Eddy Susetyo, MM",
        font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    add_formatted_paragraph(
        doc, "Anggota DPR RI A\u2013220",
        font_size=12, alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0
    )

    # Page break before documentation
    doc.add_page_break()

    # ========================================
    # DOKUMENTASI FOTO KEGIATAN
    # ========================================
    add_formatted_paragraph(
        doc, "DOKUMENTASI FOTO KEGIATAN",
        font_size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=18
    )

    for keg in KEGIATAN:
        # Tanggal
        add_formatted_paragraph(
            doc, keg["tanggal"],
            font_size=12, bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
            space_before=12, space_after=6
        )

        # Deskripsi foto
        add_formatted_paragraph(
            doc, keg["deskripsi_foto"],
            font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            space_after=6, first_line_indent=1.27
        )

        # Foto-foto
        for foto_name in keg["fotos"]:
            foto_path = os.path.join(FOTO_DIR, foto_name)
            add_photo_to_doc(doc, foto_path, width_cm=13)

        # Separator
        add_formatted_paragraph(doc, "", font_size=6, space_after=0)

    # ========================================
    # SAVE
    # ========================================
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\n[OK] Laporan berhasil dibuat: {OUTPUT_FILE}")
    print(f"     Jumlah kegiatan: {len(KEGIATAN)}")
    print(f"     Jumlah foto: {sum(len(k['fotos']) for k in KEGIATAN)}")


if __name__ == "__main__":
    create_document()
